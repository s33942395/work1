"""
分析 docx 文件內容 - 問卷描述性統計
"""
from docx import Document
import re

def analyze_docx(docx_path):
    """讀取並分析 docx 文件"""
    try:
        doc = Document(docx_path)
        
        print(f"文件總段落數: {len(doc.paragraphs)}")
        print(f"文件總表格數: {len(doc.tables)}")
        print("="*80)
        
        # 提取所有段落
        print("\n【段落內容】\n")
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:  # 只顯示非空段落
                style = para.style.name if para.style else "Normal"
                print(f"[段落 {i}] ({style})")
                print(f"{text}")
                print("-"*60)
        
        # 提取所有表格
        print("\n\n【表格內容】\n")
        for table_idx, table in enumerate(doc.tables, 1):
            print(f"\n=== 表格 {table_idx} ===")
            print(f"行數: {len(table.rows)}, 列數: {len(table.columns)}")
            
            # 顯示表格內容
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                print(f"Row {row_idx + 1}: {' | '.join(row_data)}")
            
            print("-"*80)
        
        # 分析結構
        print("\n\n【結構分析】\n")
        
        # 找出標題
        titles = []
        for para in doc.paragraphs:
            if para.style and ('Heading' in para.style.name or '標題' in para.style.name):
                titles.append((para.style.name, para.text.strip()))
        
        if titles:
            print("檢測到的標題:")
            for style, text in titles:
                print(f"  [{style}] {text}")
        
        # 找出可能的圖表標題
        print("\n可能的圖表標題/說明:")
        patterns = [
            r'圖\s*\d+',
            r'表\s*\d+',
            r'Figure\s*\d+',
            r'Table\s*\d+',
            r'圖表\s*\d+',
        ]
        
        for para in doc.paragraphs:
            text = para.text.strip()
            for pattern in patterns:
                if re.search(pattern, text):
                    print(f"  → {text}")
                    break
        
        # 統計關鍵字
        print("\n\n【統計關鍵字出現次數】")
        full_text = "\n".join([para.text for para in doc.paragraphs])
        
        keywords = [
            '平均數', '標準差', '中位數', '眾數', '百分比',
            '次數', '人數', '比率', '分佈', '統計',
            '問卷', '填答', '受訪者', '樣本',
            '公司方', '投資方', '階段'
        ]
        
        for keyword in keywords:
            count = full_text.count(keyword)
            if count > 0:
                print(f"  {keyword}: {count} 次")
        
        return doc
        
    except Exception as e:
        print(f"讀取檔案時發生錯誤: {e}")
        return None

if __name__ == "__main__":
    docx_path = "/workspaces/work1/問卷描述性統計111２.docx"
    
    print("開始分析 docx 文件...")
    print("="*80)
    
    doc = analyze_docx(docx_path)
    
    if doc:
        print("\n\n✅ 分析完成!")
    else:
        print("\n\n❌ 分析失敗")
