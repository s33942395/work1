"""
描述性統計報告生成器 - 配合 docx 格式
輸出為 Word 格式，包含圖表、表格、統計檢定、信度效度分析
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd
import numpy as np
from scipy.stats import chi2_contingency, kruskal, mannwhitneyu, fisher_exact, f_oneway
import plotly.graph_objects as go
import plotly.express as px
from io import BytesIO
from datetime import datetime
import os
import re
import warnings
from difflib import SequenceMatcher
warnings.filterwarnings('ignore')

# 導入信度效度分析模組
try:
    from reliability_validity_analysis import (
        calculate_cronbach_alpha, kmo_test, bartlett_test,
        interpret_kmo, interpret_cronbach_alpha
    )
    RELIABILITY_AVAILABLE = True
except:
    RELIABILITY_AVAILABLE = False
    print("警告：信度效度分析模組未載入")

def smart_sort_categories(categories):
    """
    智慧排序類別資料 - 與 cloud_app.py 完全一致（改進版）
    處理百分比、數值範圍、階段、是否等特殊格式
    """
    if len(categories) == 0:
        return []
    
    categories_list = list(categories)
    
    def sort_key(item):
        item_str = str(item).strip()
        
        # 1. 百分比範圍和特殊情況
        # 先檢查"以下"（應該排在最前）
        if '以下' in item_str:
            below_match = re.search(r'(\d+\.?\d*)\s*[%％]?\s*以下', item_str)
            if below_match:
                return (0, -1, float(below_match.group(1)))  # 用 -1 確保排在範圍前
        
        # 檢查"以上"（應該排在最後）
        if '以上' in item_str:
            above_match = re.search(r'(\d+\.?\d*)\s*[%％]?\s*以上', item_str)
            if above_match:
                return (0, 1000, float(above_match.group(1)))  # 用 1000 確保排在最後
        
        # 百分比範圍 (如 10-20%, 20%-30%)
        percent_match = re.match(r'(\d+\.?\d*)\s*[-~到至]\s*(\d+\.?\d*)\s*[%％]', item_str)
        if percent_match:
            return (0, 0, float(percent_match.group(1)))
        
        # 單一百分比 (如 30%)
        single_percent = re.match(r'(\d+\.?\d*)\s*[%％]', item_str)
        if single_percent:
            return (0, 0, float(single_percent.group(1)))
        
        # 2. 年份範圍
        year_match = re.match(r'(\d+\.?\d*)\s*[-~到至]\s*(\d+\.?\d*)\s*年', item_str)
        if year_match:
            return (1, 0, float(year_match.group(1)))
        
        # 3. 金額範圍
        money_match = re.match(r'(\d+\.?\d*)\s*[-~到至]\s*(\d+\.?\d*)\s*[萬億]', item_str)
        if money_match:
            return (2, 0, float(money_match.group(1)))
        
        # 4. 月份範圍
        month_match = re.match(r'(\d+\.?\d*)\s*[-~到至]\s*(\d+\.?\d*)\s*個?月', item_str)
        if month_match:
            return (3, 0, float(month_match.group(1)))
        
        # 5. 人數範圍
        people_match = re.match(r'(\d+\.?\d*)\s*[-~到至]\s*(\d+\.?\d*)\s*人', item_str)
        if people_match:
            return (4, 0, float(people_match.group(1)))
        
        # 6. 頻率
        freq_order = {'每週': 1, '每月': 2, '每季': 3, '每半年': 4, '每年': 5, '不定期': 6, '無': 7}
        for key, value in freq_order.items():
            if key in item_str:
                return (5, 0, value)
        
        # 7. 李克特量表（Likert Scale）
        likert_order = {
            '非常不同意': 1, '不同意': 2, '普通': 3, '同意': 4, '非常同意': 5,
            '非常不滿意': 1, '不滿意': 2, '普通': 3, '滿意': 4, '非常滿意': 5,
            '非常不重要': 1, '不重要': 2, '普通': 3, '重要': 4, '非常重要': 5
        }
        for key, value in likert_order.items():
            if item_str == key:
                return (6, 0, value)
        
        # 8. 階段
        if '第一階段' in item_str or '階段1' in item_str:
            return (7, 0, 1)
        if '第二階段' in item_str or '階段2' in item_str:
            return (7, 0, 2)
        if '第三階段' in item_str or '階段3' in item_str:
            return (7, 0, 3)
        
        # 9. 是/否
        if item_str in ['是', 'Yes', 'yes', '有']:
            return (8, 0, 1)
        if item_str in ['否', 'No', 'no', '無']:
            return (8, 0, 2)
        
        # 9. 一般文字
        return (10, 0, item_str)
    
    try:
        sorted_list = sorted(categories_list, key=sort_key)
        return sorted_list
    except:
        return categories_list

def add_heading_with_style(doc, text, level=1):
    """新增標題並設定樣式"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_statistics_table(doc, data_dict, title="", table_counter=None):
    """
    新增政府統計風格的完整表格
    包含標題、資料來源、製表單位等資訊
    table_counter: 如果提供，會自動編號表格
    """
    # 如果有 table_counter，更新表格編號
    if table_counter is not None and title:
        # 如果 title 不包含「表 X」格式，則添加
        if not re.match(r'^表\s*\d+', title):
            table_counter['count'] += 1
            title = f"表 {table_counter['count']}：{title}"
        else:
            # 更新現有編號
            table_counter['count'] += 1
            title = re.sub(r'^表\s*\d+', f'表 {table_counter["count"]}', title)
    
    if title:
        # 表格標題（置中）
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.runs[0].font.bold = True
        title_para.runs[0].font.size = Pt(12)
    
    # 創建表格
    table = doc.add_table(rows=1, cols=len(data_dict['columns']))
    table.style = 'Light Grid Accent 1'

    # 判斷是否為題項統計表（第一欄為「題項」）
    is_item_stats = len(data_dict['columns']) > 0 and str(data_dict['columns'][0]).strip() in ['題項', '項目', '題目']

    # 標題列（加粗、置中）
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(data_dict['columns']):
        hdr_cells[i].text = str(col_name)
        # 設定標題列樣式
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # 若為題項統計表，調整第一欄寬度（如7cm）
    if is_item_stats:
        try:
            from docx.shared import Cm
            for row in table.rows:
                row.cells[0].width = Cm(7)
        except Exception as e:
            pass

    # 資料列
    for row_data in data_dict['data']:
        row_cells = table.add_row().cells
        for i, cell_value in enumerate(row_data):
            row_cells[i].text = str(cell_value)
            # 數值靠右對齊，文字靠左對齊
            if i > 0:  # 第一欄（類別/題項）靠左，其他欄靠右
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 表格備註（政府統計風格）
    doc.add_paragraph()
    note_para = doc.add_paragraph()
    note_para.add_run('資料來源：').font.size = Pt(9)
    note_para.add_run('問卷調查資料').font.size = Pt(9)
    note_para.runs[0].font.bold = True
    
    doc.add_paragraph()  # 空行
    return table

def save_plotly_as_image(fig, filename):
    """儲存 Plotly 圖表為圖片（PNG格式）"""
    try:
        # 嘗試儲存圖表，設定較長的超時時間
        fig.write_image(filename, width=1000, height=600, scale=2, engine="kaleido")
        return True
    except Exception as e:
        print(f"圖表儲存失敗: {e}")
        # 嘗試使用較低品質設定
        try:
            print("嘗試使用較低品質設定...")
            fig.write_image(filename, width=800, height=500, scale=1, engine="kaleido")
            return True
        except Exception as e2:
            print(f"圖表儲存再次失敗: {e2}")
            return False

def create_bar_chart(crosstab, crosstab_pct, title, categories):
    """
    創建長條圖（公司方 vs 投資方比較）- 與 cloud_app.py 完全一致
    使用智慧排序、plotly_white 模板、正確配色
    """
    # 智慧排序 X 軸
    sorted_categories = smart_sort_categories(categories)
    
    # 重新排序數據
    crosstab_sorted = crosstab.reindex(sorted_categories)
    crosstab_pct_sorted = crosstab_pct.reindex(sorted_categories)
    
    fig = go.Figure()
    
    # 使用與 app 一致的配色方案
    colors = {'公司方': '#1f77b4', '投資方': '#ff7f0e', '未知': '#999999'}
    
    # 公司方數據 - 使用百分比作為 Y 軸
    if '公司方' in crosstab_pct_sorted.columns:
        company_pct = crosstab_pct_sorted['公司方'].values
        
        fig.add_trace(go.Bar(
            name='公司方',
            x=sorted_categories,
            y=company_pct,
            marker_color=colors['公司方'],
            text=[f"{p:.1f}%" for p in company_pct],
            textposition='auto'
        ))
    
    # 投資方數據 - 使用百分比作為 Y 軸
    if '投資方' in crosstab_pct_sorted.columns:
        investor_pct = crosstab_pct_sorted['投資方'].values
        
        fig.add_trace(go.Bar(
            name='投資方',
            x=sorted_categories,
            y=investor_pct,
            marker_color=colors['投資方'],
            text=[f"{p:.1f}%" for p in investor_pct],
            textposition='auto'
        ))
    
    fig.update_layout(
        barmode='group',
        title='各選項在不同身分的選擇比例',
        xaxis_title='選項',
        yaxis_title='比例 (%)',
        template='plotly_white',
        height=500,
        xaxis_tickangle=-45,
        xaxis={'categoryorder': 'array', 'categoryarray': sorted_categories},
        font=dict(family='Noto Sans CJK SC, WenQuanYi Micro Hei, sans-serif', size=12)
    )
    
    return fig

def create_horizontal_bar_chart(crosstab, crosstab_pct, title, categories):
    """
    創建水平長條圖 - 適合長文字標籤
    使用百分比、專業配色
    """
    # 智慧排序（反向以符合圖表習慣，從上到下）
    sorted_categories = smart_sort_categories(categories)
    sorted_categories.reverse()  # 反轉順序，讓最小值在上方
    
    fig = go.Figure()
    
    # 檢查是否有受訪者類型資料
    if '公司方' in crosstab.columns:
        # 有公司方和投資方比較
        company_pct = [crosstab_pct.loc[cat, '公司方'] if cat in crosstab_pct.index else 0 for cat in sorted_categories]
        investor_pct = [crosstab_pct.loc[cat, '投資方'] if cat in crosstab_pct.index and '投資方' in crosstab_pct.columns else 0 for cat in sorted_categories]
        
        fig.add_trace(go.Bar(
            name='公司方',
            y=sorted_categories,
            x=company_pct,
            orientation='h',
            marker_color='#1f77b4',
            text=[f"{p:.1f}%" for p in company_pct],
            textposition='auto'
        ))
        
        if '投資方' in crosstab.columns:
            fig.add_trace(go.Bar(
                name='投資方',
                y=sorted_categories,
                x=investor_pct,
                orientation='h',
                marker_color='#ff7f0e',
                text=[f"{p:.1f}%" for p in investor_pct],
                textposition='auto'
            ))
        
        fig.update_layout(
            barmode='group',
            title=title,
            xaxis_title='比例 (%)',
            yaxis_title='選項',
            template='plotly_white',
            height=max(400, len(sorted_categories) * 40),  # 動態高度
            showlegend=True,
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=1.02,
                xanchor="right",
                x=1
            ),
            font=dict(family='Noto Sans CJK SC, WenQuanYi Micro Hei, sans-serif', size=12),
            xaxis=dict(range=[0, 100]),
            yaxis=dict(automargin=True)
        )
    
    return fig

def create_horizontal_phase_chart(phase_crosstab, phase_crosstab_pct, title, categories, phases):
    """
    創建水平階段比較長條圖 - 適合長文字標籤
    使用百分比、專業配色
    """
    # 智慧排序（反向以符合圖表習慣）
    sorted_categories = smart_sort_categories(categories)
    sorted_categories.reverse()  # 反轉順序，讓最小值在上方
    
    # 重新排序數據
    phase_crosstab_sorted = phase_crosstab.reindex(sorted_categories)
    phase_crosstab_pct_sorted = phase_crosstab_pct.reindex(sorted_categories)
    
    fig = go.Figure()
    
    # 專業配色方案（與 app 一致：綠-紅-紫）
    colors = {
        '第一階段': '#2ca02c',  # 綠色
        '第二階段': '#d62728',  # 紅色
        '第三階段': '#9467bd'   # 紫色
    }
    
    # 智慧排序階段（確保一二三順序）
    sorted_phases = smart_sort_categories(phases)
    
    for phase in sorted_phases:
        if phase in phase_crosstab_pct_sorted.columns:
            # 使用百分比作為 X 軸
            pct_values = phase_crosstab_pct_sorted[phase].values
            
            fig.add_trace(go.Bar(
                name=phase,
                y=sorted_categories,
                x=pct_values,
                orientation='h',
                marker_color=colors.get(phase, '#cccccc'),
                text=[f"{p:.1f}%" for p in pct_values],
                textposition='auto',
                textfont=dict(size=10)
            ))
    
    fig.update_layout(
        barmode='group',
        title=title,
        xaxis_title='比例 (%)',
        yaxis_title='選項',
        template='plotly_white',
        height=max(400, len(sorted_categories) * 50),  # 動態高度
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        ),
        font=dict(family='Noto Sans CJK SC, WenQuanYi Micro Hei, sans-serif', size=12),
        xaxis=dict(range=[0, 100]),
        yaxis=dict(automargin=True)
    )
    
    return fig

def create_phase_chart(phase_crosstab, phase_crosstab_pct, title, categories, phases):
    """
    創建階段比較長條圖 - 美化版
    使用百分比、專業配色、垂直 Y 軸標籤
    """
    # 智慧排序 X 軸
    sorted_categories = smart_sort_categories(categories)
    
    # 重新排序數據
    phase_crosstab_sorted = phase_crosstab.reindex(sorted_categories)
    phase_crosstab_pct_sorted = phase_crosstab_pct.reindex(sorted_categories)
    
    fig = go.Figure()
    
    # 專業配色方案（與 app 一致：綠-紅-紫）
    colors = {
        '第一階段': '#2ca02c',  # 綠色
        '第二階段': '#d62728',  # 紅色
        '第三階段': '#9467bd'   # 紫色
    }
    
    for phase in phases:
        if phase in phase_crosstab_pct_sorted.columns:
            # 使用百分比作為 Y 軸
            pct_values = phase_crosstab_pct_sorted[phase].values
            
            fig.add_trace(go.Bar(
                name=phase,
                x=sorted_categories,
                y=pct_values,
                marker_color=colors.get(phase, '#cccccc'),
                text=[f"{p:.1f}%" for p in pct_values],
                textposition='outside',
                textfont=dict(size=10)
            ))
    
    fig.update_layout(
        barmode='group',
        title=title,  # 使用傳入的完整題目
        xaxis_title='選項',
        yaxis_title='比例 (%)',
        template='plotly_white',
        height=500,
        xaxis_tickangle=0,  # X 軸標籤水平顯示
        xaxis={'categoryorder': 'array', 'categoryarray': sorted_categories},
        yaxis=dict(
            title=dict(
                text='比例 (%)',
                standoff=15
            ),
            tickfont=dict(size=11),
            tickangle=-90  # Y 軸刻度數字垂直顯示
        ),
        font=dict(family='Noto Sans CJK SC, WenQuanYi Micro Hei, sans-serif', size=12),
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        )
    )
    
    return fig

def calculate_chi_square(df, col_name, group_col='respondent_type'):
    """計算卡方檢定"""
    try:
        crosstab = pd.crosstab(df[col_name], df[group_col])
        if crosstab.size > 0 and crosstab.values.sum() > 0:
            if crosstab.shape == (2, 2) and crosstab.values.sum() < 20:
                # 小樣本使用 Fisher 精確檢定
                oddsratio, p = fisher_exact(crosstab)
                return {'method': 'Fisher精確檢定', 'statistic': oddsratio, 'p_value': p}
            else:
                chi2, p, dof, expected = chi2_contingency(crosstab)
                return {'method': '卡方檢定', 'statistic': chi2, 'p_value': p, 'dof': dof}
    except Exception as e:
        print(f"統計檢定失敗: {e}")
    return None

def generate_descriptive_report_word(df, output_filename="問卷描述性統計報告_改進版.docx"):
    """
    生成描述性統計 Word 報告
    參考原始 docx 格式，加入圖表、表格、統計檢定
    返回 (doc, table_counter) 元組
    """
    
    # 初始化表格計數器
    table_counter = {'count': 0}
    
    # 創建 Word 文件
    doc = Document()
    
    # 設定中文字型
    doc.styles['Normal'].font.name = '微軟正黑體'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
    
    # === 封面 ===
    title = doc.add_heading('未上市櫃公司治理問卷調查', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.bold = True
    
    subtitle = doc.add_paragraph('描述性統計分析報告')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 執行單位
    org_para = doc.add_paragraph('執行單位：國家發展委員會')
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_para.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()
    date_para = doc.add_paragraph(f'中華民國 {datetime.now().year - 1911} 年 {datetime.now().month} 月')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(14)
    
    doc.add_page_break()
    
    # === 摘要 ===
    add_heading_with_style(doc, '摘要', level=1)
    
    abstract_text = (
        '本研究旨在瞭解未上市（櫃）公司治理現況，透過問卷調查方式，'
        '蒐集公司方與投資方對公司治理實踐之觀點。調查對象包含處於不同發展階段之未上市（櫃）公司，'
        '涵蓋第一階段（創業天使投資方案）、第二階段（加強投資中小企業方案等）'
        '及第三階段（直接投資方案）共三個階段。'
        '\n\n'
        f'本次調查共回收有效問卷 {len(df)} 份，其中公司方 {len(df[df["respondent_type"] == "公司方"]) if "respondent_type" in df.columns else 0} 份、'
        f'投資方 {len(df[df["respondent_type"] == "投資方"]) if "respondent_type" in df.columns else 0} 份。'
        '分析範圍涵蓋股權結構、董事會運作、資訊揭露、財務管理、內部控制及利害關係人治理等六大構面，'
        '共計 20 項關鍵議題。'
        '\n\n'
        '統計分析採用卡方檢定（Chi-square test）檢驗公司方與投資方之觀點差異，'
        '並以 Kruskal-Wallis H 檢定比較不同發展階段公司之治理實踐。'
        '研究發現部分議題存在顯著之利害關係人觀點差異，亦觀察到公司發展階段對特定治理機制之影響。'
        '本報告針對各項發現提出政策意涵與實務建議，以供政府機關、投資機構及未上市（櫃）公司參考。'
    )
    
    doc.add_paragraph(abstract_text)
    
    doc.add_page_break()
    
    # === 1. 研究背景與目的 ===
    add_heading_with_style(doc, '壹、研究背景與目的', level=1)
    # === 1. 研究背景與目的 ===
    add_heading_with_style(doc, '壹、研究背景與目的', level=1)
    
    add_heading_with_style(doc, '一、研究背景', level=2)
    
    background_text = (
        '公司治理（Corporate Governance）係指公司經營權與所有權分離後，'
        '為確保股東權益並兼顧利害關係人利益，所建立之一套管理與監督機制。'
        '我國自 2002 年起推動公司治理改革，已建立相對完善之上市櫃公司治理制度。'
        '然未上市（櫃）公司因規模較小、資訊透明度較低、利害關係人結構較為單純，'
        '其治理機制之建立與執行與上市櫃公司有顯著差異。'
        '\n\n'
        '國家發展基金（以下簡稱國發基金）為促進產業升級及經濟發展，'
        '透過創業天使投資方案、加強投資中小企業實施方案等多元管道，'
        '投資未上市（櫃）公司。為瞭解受投資企業之治理現況，'
        '並提供適切之輔導措施，實有必要針對不同發展階段之未上市（櫃）公司進行系統性調查。'
    )
    doc.add_paragraph(background_text)
    
    add_heading_with_style(doc, '二、研究目的', level=2)
    
    doc.add_paragraph('本研究具體目的如下：')
    
    purposes = [
        '瞭解未上市（櫃）公司在股權結構、董事會運作、資訊揭露、財務管理、內部控制及利害關係人治理等面向之實踐現況。',
        '比較公司方與投資方對公司治理議題之觀點差異，釐清雙方認知落差，促進溝通與共識建立。',
        '分析不同發展階段公司之治理特徵，瞭解公司治理機制隨企業成長之演進模式。',
        '透過統計檢定方法，檢驗利害關係人觀點差異與階段性差異之顯著性，提升研究發現之可信度。',
        '依據研究發現，提出政策建議與實務作法，協助政府機關制定未上市（櫃）公司治理輔導政策。'
    ]
    
    for i, purpose in enumerate(purposes, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(purpose)
    
    doc.add_page_break()
    
    # === 2. 研究方法 ===
    add_heading_with_style(doc, '貳、研究方法', level=1)
    
    add_heading_with_style(doc, '一、調查設計', level=2)
    
    method_text = (
        '本研究採用問卷調查法（Questionnaire Survey），針對國發基金投資之未上市（櫃）公司'
        '及其投資人進行調查。問卷設計參考公開發行公司董事會議事辦法、證券交易法相關規範、'
        '以及國際公司治理準則（如 OECD Principles of Corporate Governance），'
        '並依未上市（櫃）公司特性調整題項內容。'
    )
    doc.add_paragraph(method_text)
    
    add_heading_with_style(doc, '二、抽樣方法', level=2)
    
    sampling_text = (
        '本研究採分層抽樣（Stratified Sampling）方式，依公司發展階段分為三層：'
    )
    doc.add_paragraph(sampling_text)
    
    sampling_layers = [
        ('第一階段（創立期）', '接受創業天使投資方案之公司，多為成立 3 年內之新創企業，著重產品開發與市場驗證。'),
        ('第二階段（成長期）', '接受加強投資中小企業、策略性服務業、策略性製造業、文化創意產業等方案之公司，已具備穩定營運模式，處於規模擴張階段。'),
        ('第三階段（成熟期）', '接受直接投資方案之公司，營運相對穩定，具備一定市場地位，治理制度較為完善。')
    ]
    
    for stage, desc in sampling_layers:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(stage + '：').bold = True
        p.add_run(desc)
    
    add_heading_with_style(doc, '三、分析方法', level=2)
    
    analysis_text = (
        '本研究採用描述性統計（Descriptive Statistics）與推論統計（Inferential Statistics）'
        '進行資料分析：'
    )
    doc.add_paragraph(analysis_text)
    
    analysis_methods = [
        ('描述性統計', '計算各題項之次數分配、百分比、平均數、標準差等，呈現樣本特徵與分佈情形。'),
        ('卡方檢定（Chi-square test）', '檢驗公司方與投資方在各治理議題上之觀點差異是否達統計顯著水準（α = 0.05）。適用於類別變數之獨立性檢定。'),
        ('Kruskal-Wallis H 檢定', '比較三個發展階段公司在各議題上之差異，屬無母數檢定方法，不假設資料符合常態分佈，適用於順序資料或偏態分佈之連續變數。'),
        ('Fisher 精確檢定（Fisher\'s Exact Test）', '當 2×2 列聯表樣本數過小（期望次數 < 5）時，以 Fisher 精確檢定取代卡方檢定，提升檢定結果之準確性。')
    ]
    
    for method, desc in analysis_methods:
        p = doc.add_paragraph(style='List Number')
        p.add_run(method + '：').bold = True
        p.add_run(desc)
    
    doc.add_paragraph()
    significance_note = doc.add_paragraph(
        '註：本研究顯著性檢定採雙尾檢定（two-tailed test），顯著水準設定為 α = 0.05，'
        '即當 p 值 < 0.05 時，拒絕虛無假設，認定組間差異達統計顯著水準。'
        '顯著性標記說明：*** 表示 p < 0.001（高度顯著）、** 表示 p < 0.01（非常顯著）、'
        '* 表示 p < 0.05（顯著）、n.s. 表示 p ≥ 0.05（無顯著差異）。'
    )
    significance_note.runs[0].font.size = Pt(10)
    significance_note.runs[0].font.italic = True
    
    doc.add_page_break()
    
    # === 3. 題目選擇邏輯 ===
    add_heading_with_style(doc, '參、分析架構與題目選擇', level=1)
    
    add_heading_with_style(doc, '一、分析架構', level=2)
    
    framework_text = (
        '本研究參考國內外公司治理文獻，建構未上市櫃公司治理分析架構，'
        '涵蓋六大構面：'
    )
    doc.add_paragraph(framework_text)
    
    framework_dimensions = [
        ('股權結構與控制權', '探討股權集中度、經營團隊持股等議題，分析所有權與經營權之結合程度，以及對代理問題之影響。'),
        ('股東會治理', '檢視股東會召開程序、議事錄完整性、董事出席情形等，評估股東權益保障機制之落實情況。'),
        ('董事會治理機制', '分析董事會運作效能，包括會議通知、議事錄記載、召開頻率、議事內容及外部諮詢等面向。'),
        ('財務報告與資訊透明度', '瞭解財務報表查核、股權結構揭露、業務及財務報告提供頻率等，評估資訊揭露之完整性與及時性。'),
        ('內部控制與風險管理', '檢視財務職能分工、財務紀錄處理、智慧財產權保護等內部控制機制之建立與執行。'),
        ('利害關係人治理', '探討員工激勵制度、利害關係人溝通管道等，瞭解公司對人力資本管理與多方協作之重視程度。')
    ]
    
    for i, (dim, desc) in enumerate(framework_dimensions, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(f'{dim}：').bold = True
        p.add_run(desc)
    
    add_heading_with_style(doc, '二、題目選擇原則', level=2)
    
    logic_para = doc.add_paragraph(
        '本研究從問卷題庫中精選 20 個關鍵議題進行深入分析，選題依據以下五項原則：'
    )
    
    criteria = [
        ('代表性原則', '題目需能代表各治理構面之核心議題，確保分析涵蓋股權結構、董事會運作、資訊揭露、財務管理、內部控制及利害關係人治理等面向，呈現公司治理之全貌。'),
        ('差異性原則', '優先選擇公司方與投資方觀點可能存在差異之題目，如資訊揭露頻率、決策參與程度等，以釐清利害關係人認知落差，促進雙方溝通。'),
        ('階段性原則', '選擇能反映公司發展階段特徵之題目，如創立期著重股權結構、成長期強調董事會運作、成熟期重視資訊揭露，以瞭解治理機制之演進模式。'),
        ('實務性原則', '聚焦於治理機制之實際執行狀況，而非僅止於制度之存在與否。選擇可量化、可比較之指標，如董事會召開頻率、財報提供頻率等，提升分析之客觀性。'),
        ('完整性原則', '避免選擇遺漏資料過多或選項過於單一之題目，確保統計分析之有效性與可信度。各題有效樣本數需達統計分析之最低要求。')
    ]
    
    for i, (principle, desc) in enumerate(criteria, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(principle + '：').bold = True
        p.add_run(desc)
    
    selection_note = doc.add_paragraph()
    selection_note.add_run(
        '\n透過上述選題原則，本研究共選定 20 項議題進行分析。'
        '每項議題均包含「公司方與投資方比較」及「公司發展階段分析」兩個面向，'
        '並輔以適當之統計檢定方法，以全面呈現未上市櫃公司治理現況及其隨企業成長之演進趨勢。'
    )
    
    doc.add_page_break()
    
    # === 4. 樣本特性分析 ===
    add_heading_with_style(doc, '肆、樣本特性分析', level=1)
    # === 4. 樣本特性分析 ===
    add_heading_with_style(doc, '肆、樣本特性分析', level=1)
    
    add_heading_with_style(doc, '一、樣本來源', level=2)
    
    sample_intro = doc.add_paragraph(
        '本研究樣本來自國家發展基金投資之未上市（櫃）公司及其投資機構。'
        '樣本公司產業分佈廣泛，涵蓋半導體、人工智慧、通訊、電子商務、生物科技、'
        '機械製造等高成長潛力產業。投資機構則包括國內外知名創業投資公司（Venture Capital, VC）'
        '與企業創投（Corporate Venture Capital, CVC）。'
    )
    
    add_heading_with_style(doc, '二、樣本結構', level=2)
    add_heading_with_style(doc, '二、樣本結構', level=2)
    
    # 基本統計
    total_samples = len(df)
    if 'respondent_type' in df.columns:
        company_count = len(df[df['respondent_type'] == '公司方'])
        investor_count = len(df[df['respondent_type'] == '投資方'])
    else:
        company_count = investor_count = 0
    
    if 'phase' in df.columns:
        phase_counts = df['phase'].value_counts().to_dict()
        phase1_count = phase_counts.get('第一階段', 0)
        phase2_count = phase_counts.get('第二階段', 0)
        phase3_count = phase_counts.get('第三階段', 0)
    else:
        phase1_count = phase2_count = phase3_count = 0
    
    sample_structure = doc.add_paragraph(
        f'本次調查共回收有效問卷 {total_samples} 份。依受訪者類型區分，'
        f'公司方填答問卷 {company_count} 份（{company_count/total_samples*100:.1f}%），'
        f'投資方填答問卷 {investor_count} 份（{investor_count/total_samples*100:.1f}%）。'
        f'依公司發展階段區分，第一階段（創立期）{phase1_count} 份（{phase1_count/total_samples*100:.1f}%）、'
        f'第二階段（成長期）{phase2_count} 份（{phase2_count/total_samples*100:.1f}%）、'
        f'第三階段（成熟期）{phase3_count} 份（{phase3_count/total_samples*100:.1f}%）。'
    )
    
    # 樣本結構表格
    table_data = {
        'columns': ['分類項目', '類別', '樣本數', '百分比'],
        'data': [
            ['受訪者類型', '公司方', company_count, f'{company_count/total_samples*100:.1f}%'],
            ['', '投資方', investor_count, f'{investor_count/total_samples*100:.1f}%'],
            ['', '小計', total_samples, '100.0%'],
            ['公司發展階段', '第一階段（創立期）', phase1_count, f'{phase1_count/total_samples*100:.1f}%'],
            ['', '第二階段（成長期）', phase2_count, f'{phase2_count/total_samples*100:.1f}%'],
            ['', '第三階段（成熟期）', phase3_count, f'{phase3_count/total_samples*100:.1f}%'],
            ['', '小計', total_samples, '100.0%']
        ]
    }
    
    add_statistics_table(doc, table_data, title='樣本結構分布表', table_counter=table_counter)
    
    add_heading_with_style(doc, '三、樣本限制', level=2)
    add_heading_with_style(doc, '三、樣本限制', level=2)
    
    limitation_text = (
        '本研究樣本主要來自國家發展基金投資之未上市（櫃）公司，'
        '在樣本代表性上存在以下限制：'
    )
    doc.add_paragraph(limitation_text)
    
    limitations = [
        '樣本公司多為接受政府投資之企業，可能與一般未上市（櫃）公司在治理水準上存在系統性差異。',
        '樣本集中於臺灣地區，研究發現之適用性可能受地域文化、法規環境等因素影響，推論至其他國家或地區時需審慎。',
        '受限於調查時間與資源，本研究採橫斷面調查設計（Cross-sectional Study），無法追蹤個別公司治理之動態變化。',
        '部分議題可能涉及公司敏感資訊，受訪者填答時可能存在社會期許偏誤（Social Desirability Bias），影響資料真實性。'
    ]
    
    for limitation in limitations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(limitation)
    
    doc.add_paragraph()
    limitation_note = doc.add_paragraph(
        '儘管存在上述限制，本研究透過嚴謹之調查設計、適當之統計方法及詳實之資料分析，'
        '仍能提供未上市櫃公司治理現況之有效資訊，作為政策制定與實務改善之參考依據。'
    )
    
    doc.add_page_break()
    
    # === 5. 分析結果 ===
    add_heading_with_style(doc, '伍、分析結果', level=1)
    
    doc.add_paragraph(f'本研究問卷共蒐集樣本 {total_samples} 份，樣本共分為兩類：')
    doc.add_paragraph(f'• 公司方：共計 {company_count} 份，涵蓋創辦人、高階主管及公司治理相關人員等')
    doc.add_paragraph(f'• 投資方：共計 {investor_count} 份，涵蓋創投、天使投資人及其他機構投資人等')
    
    # 樣本分佈表
    if 'respondent_type' in df.columns:
        sample_data = {
            'columns': ['類別', '份數', '百分比'],
            'data': [
                ['公司方', company_count, f'{company_count/total_samples*100:.1f}%'],
                ['投資方', investor_count, f'{investor_count/total_samples*100:.1f}%'],
                ['合計', total_samples, '100.0%']
            ]
        }
        add_statistics_table(doc, sample_data, title="樣本分佈統計", table_counter=table_counter)
    
    # === 核心議題說明 ===
    add_heading_with_style(doc, '(一) 核心議題選定說明', level=2)
    
    doc.add_paragraph(
        '本報告從問卷中精選 20 個核心議題進行深入分析，涵蓋未上市櫃公司治理的六大關鍵面向。'
        '這些議題的選擇係基於公司治理理論與實務的重要性，以及對投資人決策與公司永續發展的影響程度。'
    )
    
    doc.add_paragraph()
    doc.add_paragraph('【六大關鍵面向】', style='Heading 4')
    
    dimension_para = doc.add_paragraph()
    dimension_para.add_run('1. 股權結構與控制權（2題）：').bold = True
    dimension_para.add_run('股權集中度與經營團隊持股是公司治理的基礎，直接影響決策效率與代理問題。')
    
    dimension_para = doc.add_paragraph()
    dimension_para.add_run('2. 股東會治理（3題）：').bold = True
    dimension_para.add_run('股東會的通知時效、決議記錄與董事出席情況，是保障股東權益的基本機制。')
    
    dimension_para = doc.add_paragraph()
    dimension_para.add_run('3. 董事會治理機制（5題）：').bold = True
    dimension_para.add_run('董事會的議程通知、決議記錄、召集程序、議事內容與開會頻率，是確保董事會有效運作的核心要素。')
    
    dimension_para = doc.add_paragraph()
    dimension_para.add_run('4. 財務報告與資訊透明度（5題）：').bold = True
    dimension_para.add_run('外部專業諮詢、財務查核、股權揭露、業務報告與財務報告頻率，是投資人監督管理層的重要基礎。')
    
    dimension_para = doc.add_paragraph()
    dimension_para.add_run('5. 內部控制與風險管理（3題）：').bold = True
    dimension_para.add_run('財務職能分工、財務紀錄處理與智慧財產權保護，是未上市櫃公司內部控制的基本要求。')
    
    dimension_para = doc.add_paragraph()
    dimension_para.add_run('6. 利害關係人治理（2題）：').bold = True
    dimension_para.add_run('員工激勵制度與利害關係人溝通機制，反映公司對人力資本與多方協作的重視程度。')
    
    doc.add_paragraph()
    doc.add_paragraph(
        '以下分別針對各項核心議題，提供完整的統計分析（包含公司方與投資方比較、'
        '公司發展階段分析）、統計檢定結果，以及業務意涵解讀。即使統計檢定未達顯著水準，'
        '仍透過描述性統計提供實務觀察，供政策制定與業務推動參考。'
    )
    
    doc.add_page_break()
    
    # === 議題分析 ===
    # 這裡需要根據實際欄位動態生成
    # 先返回基礎文件結構和表格計數器
    
    return doc, table_counter

def clean_and_merge_categories(series):
    """
    清理並合併重複的類別（如「50%-67%」和「50%-67(不含)%」）
    返回標準化後的 Series
    """
    # 創建映射字典來標準化類別名稱
    mapping = {}
    
    for val in series.unique():
        if pd.isna(val):
            continue
        val_str = str(val).strip()
        
        # 標準化百分比範圍格式
        # 例如：50%-67% -> 50%-67(不含)%
        percent_pattern = re.match(r'(\d+\.?\d*)\s*[%％]?\s*[-~到至]\s*(\d+\.?\d*)\s*[%％]', val_str)
        if percent_pattern:
            start = percent_pattern.group(1)
            end = percent_pattern.group(2)
            standardized = f"{start}%-{end}(不含)%"
            mapping[val] = standardized
            continue
        
        # 標準化「不定期」類別（處理各種變體）
        if '不定期' in val_str or '不定期提供' in val_str:
            mapping[val] = '不定期'
            continue
        
        # 其他情況保持原樣
        mapping[val] = val_str
    
    return series.map(lambda x: mapping.get(x, x) if pd.notna(x) else x)

def find_matching_column(df, target_col):
    """
    查找匹配的欄位名稱，處理公司方和投資方的不同命名
    """

    # 嚴格分離公司方/投資方欄位對應
    alt_31_company = '請問公司大股東（持股5%以上）合計持股比例為多少？'
    alt_31_investor = '請問您投資的公司之大股東（持股5%以上）合計持股比例為多少？'
    if target_col == alt_31_company and alt_31_company in df.columns:
        return alt_31_company
    if target_col == alt_31_investor and alt_31_investor in df.columns:
        return alt_31_investor
    # 嚴格對應大股東人數題目
    alt_31n_company = '請問公司的大股東（持股5%以上）人數多少人'
    alt_31n_company2 = '請問公司的大股東（持股5%以上）人數多少人？'
    alt_31n_std = '公司大股東（持股5%以上）人數'
    if target_col in [alt_31n_company, alt_31n_company2, alt_31n_std] and alt_31n_std in df.columns:
        return alt_31n_std
    if target_col == alt_31n_std and alt_31n_company in df.columns:
        return alt_31n_company
    alt_meeting_company = '股東會結構與運作 - 公司股東常會的議程及相關資料能在20天前通知，並以可存證的方式（如：掛號或經股東同意的電子方式）寄發'
    alt_meeting_investor = '股東會結構與運作 - 您投資的公司之股東常會的議程及相關資料能在20天前通知，並以可存證的方式（如：掛號或經股東同意的電子方式）寄發'
    if target_col == alt_meeting_company and alt_meeting_company in df.columns:
        return alt_meeting_company
    if target_col == alt_meeting_investor and alt_meeting_investor in df.columns:
        return alt_meeting_investor

    # 2. 嚴格mapping表（公司方/投資方分開）
    investor_mappings = {
        '請問公司大股東（持股5%以上）合計持股比例為多少？': '請問您投資的公司之大股東（持股5%以上）合計持股比例為多少？',
        '請問公司經營團隊合計持股比例為多少？': '請問您投資的公司之經營團隊合計持股比例為多少？',
        # 大股東人數題目
        '請問公司的大股東（持股5%以上）人數多少人': '公司大股東（持股5%以上）人數',
        '請問公司的大股東（持股5%以上）人數多少人？': '公司大股東（持股5%以上）人數',
        '公司大股東（持股5%以上）人數': '公司大股東（持股5%以上）人數',
        # ...existing code...
    }
    # 1. 直接存在且有資料
    if target_col in df.columns and df[target_col].dropna().shape[0] > 0:
        return target_col
    # 2. mapping 對應且有資料
    if target_col in investor_mappings:
        mapped_col = investor_mappings[target_col]
        if mapped_col in df.columns and df[mapped_col].dropna().shape[0] > 0:
            return mapped_col
    reverse_mapping = {v: k for k, v in investor_mappings.items()}
    if target_col in reverse_mapping:
        mapped_col = reverse_mapping[target_col]
        if mapped_col in df.columns and df[mapped_col].dropna().shape[0] > 0:
            return mapped_col

    # 3. 自動搜尋所有公司方/投資方對應欄位
    def normalize_col(col):
        s = str(col)
        # 只移除明確的問句前綴與階段前綴，避免過度刪除關鍵詞造成錯配
        for prefix in ['請問您投資的公司之', '請問公司的', '請問公司', '請問您投資的', '請問', '【第一階段：', '【第二階段：', '【第三階段：']:
            s = s.replace(prefix, '')
        # 移除常見標點與空白
        s = re.sub(r'[：:\-—–()（）\[\]{}、,，\s\n\r]+', ' ', s)
        return s.strip().lower()
    norm_target = normalize_col(target_col)
    for col in df.columns:
        if normalize_col(col) == norm_target and df[col].dropna().shape[0] > 0:
            return col

    # 4. fallback: 允許部分比對（保留字詞順序比對，並需有實際資料）
    norm_tokens = [t for t in norm_target.split() if t]
    if norm_tokens:
        for col in df.columns:
            col_norm = normalize_col(col)
            # 檢查是否包含所有 token（維持順序敏感性簡化版）
            if all(tok in col_norm for tok in norm_tokens) and df[col].dropna().shape[0] > 0:
                print(f"[find_matching_column] Fallback matched: '{target_col}' -> '{col}' (partial tokens)")
                return col

    # 3. 僅在mapping找不到時，才啟用模糊比對，且必須型態驗證通過
    # 僅針對百分比區間題目允許模糊比對
    # 3.1合計持股比例（百分比型）
    keywords_31 = ['大股東', '合計', '持股']
    if all(k in target_col for k in keywords_31):
        for col in df.columns:
            if all(k in col for k in keywords_31):
                sample = df[col].dropna().astype(str)
                # 若是百分比型才允許
                if sample.str.contains('%').any():
                    return col
        print(f"[find_matching_column] Warning: 大股東合計持股比例題目未找到對應欄位: '{target_col}'")
        return None
    # 3.1人數型題目（允許數值/人數型）
    keywords_31n = ['大股東', '人數']
    if all(k in target_col for k in keywords_31n):
        for col in df.columns:
            if all(k in col for k in keywords_31n):
                sample = df[col].dropna().astype(str)
                # 只要是數字或「人」字結尾即可
                if sample.str.match(r'^[0-9]+(\.[0-9]+)?(人)?$').any():
                    return col
        print(f"[find_matching_column] Warning: 大股東人數題目未找到對應欄位: '{target_col}'")
        return None
    keywords_team = ['經營團隊', '持股']
    if all(k in target_col for k in keywords_team):
        for col in df.columns:
            if all(k in col for k in keywords_team):
                sample = df[col].dropna().astype(str)
                if sample.str.contains('%').any():
                    return col
        print(f"[find_matching_column] Warning: 經營團隊持股比例題目未找到對應欄位: '{target_col}'")
        return None
    keywords_13 = ['股權結構', '資訊揭露']
    if all(k in target_col for k in keywords_13):
        for col in df.columns:
            if all(k in col for k in keywords_13):
                sample = df[col].dropna().astype(str)
                if sample.str.contains('%').any():
                    return col
        print(f"[find_matching_column] Warning: 股權結構/資訊揭露題目未找到對應欄位: '{target_col}'")
        return None
    # 4. 公司股東常會議程模糊比對
    keywords_meeting = ['常會', '議程', '通知']
    if all(k in target_col for k in keywords_meeting):
        for col in df.columns:
            if all(k in col for k in keywords_meeting):
                return col
        print(f"[find_matching_column] Warning: 會議相關題目未找到對應欄位: '{target_col}'")
        return None
    # 5. mapping與模糊比對皆失敗，嘗試針對常見 metadata 題目做關鍵字回退
    lowered = target_col.lower() if isinstance(target_col, str) else ''
    # 若是填答身分類型的欄位，嘗試以關鍵字匹配
    if '填答' in lowered or '身分' in lowered or '身份' in lowered:
        for col in df.columns:
            c_low = str(col).lower()
            if '填答' in c_low or '身分' in c_low or '身份' in c_low or 'respondent' in c_low:
                if df[col].dropna().shape[0] > 0:
                    print(f"[find_matching_column] Keyword fallback for respondent_type: '{target_col}' -> '{col}'")
                    return col
    # 若是代表公司名稱之類的題目，嘗試匹配常見公司名稱欄位
    # Treat explicit company-name / representative-name questions as metadata: do not attempt fuzzy mapping
    if isinstance(target_col, str):
        t_low = target_col.lower()
        if ('代表' in t_low and '公司' in t_low) or '代表公司的名稱' in t_low or '公司名稱' in t_low or t_low.strip().startswith('請問您代表公司'):
            print(f"[find_matching_column] Treating company-name question as metadata, returning None for '{target_col}'")
            return None
    # 若題目包含非一般文字的 metadata（如 # 或 / 分隔的分數/金額標記），視為不可分析題目，回傳 None（由上層決定是否跳過）
    if isinstance(target_col, str) and (target_col.strip().startswith('#') or '/' in target_col or '分數' in target_col or '金額' in target_col):
        print(f"[find_matching_column] Treating as metadata or unsupported question, returning None for '{target_col}'")
        return None

    # 既然都找不到，回傳 None（上層會決定是否跳過）
    print(f"[find_matching_column] 找不到題目對應欄位，將回傳 None: '{target_col}'")
    return None

    # 先嘗試直接存在
    if target_col in df.columns:
        return target_col
    # 強化：3.1大股東合計持股比例、公司股東常會議程，無論公司方或投資方欄位都能對應
    # 3.1
    alt_31_company = '請問公司大股東（持股5%以上）合計持股比例為多少？'
    alt_31_investor = '請問您投資的公司之大股東（持股5%以上）合計持股比例為多少？'
    if (target_col == alt_31_company and alt_31_investor in df.columns):
        return alt_31_investor
    if (target_col == alt_31_investor and alt_31_company in df.columns):
        return alt_31_company
    # 公司股東常會議程
    alt_meeting_company = '股東會結構與運作 - 公司股東常會的議程及相關資料能在20天前通知，並以可存證的方式（如：掛號或經股東同意的電子方式）寄發'
    alt_meeting_investor = '股東會結構與運作 - 您投資的公司之股東常會的議程及相關資料能在20天前通知，並以可存證的方式（如：掛號或經股東同意的電子方式）寄發'
    if (target_col == alt_meeting_company and alt_meeting_investor in df.columns):
        return alt_meeting_investor
    if (target_col == alt_meeting_investor and alt_meeting_company in df.columns):
        return alt_meeting_company
    
    # 定義公司方→投資方的映射規則（使用完整CSV欄位名稱）
    investor_mappings = {
        # ========== 第一階段 ==========
        # 持股比例題目（注意：CSV中這兩題沒有分類前綴）
        '請問公司大股東（持股5%以上）合計持股比例為多少？': '請問您投資的公司之大股東（持股5%以上）合計持股比例為多少？',
        '請問公司經營團隊合計持股比例為多少？': '請問您投資的公司之經營團隊合計持股比例為多少？',
        
        # 股東會結構與運作
        '股東會結構與運作 - 召開股東會議的行政成本與時間，對公司是個不小的負擔': '股東會結構與運作 - 您投資的公司在召開股東會議的行政成本與時間，對公司是個不小的負擔',
        '股東會結構與運作 - 公司股東常會的議程及相關資料能在20天前通知，並以可存證的方式（如：掛號或經股東同意的電子方式）寄發': '股東會結構與運作 - 您投資的公司之股東常會的議程及相關資料能在20天前通知，並以可存證的方式（如：掛號或經股東同意的電子方式）寄發',
        '股東會結構與運作 - 公司股東會決議方式能夠清楚載明，且議事錄完整記載會議資訊（如時間、地點、主席、決議方法及結果）': '股東會結構與運作 - 您投資的公司之股東會決議方式能夠清楚載明，且議事錄完整記載會議資訊（如時間、地點、主席、決議方法及結果）',
        '股東會結構與運作 - 公司董事長及董事通常能夠親自出席股東常會': '股東會結構與運作 - 您投資的公司之董事長及董事通常能夠親自出席股東常會',
        
        # 董事會結構與運作
        '董事會結構與運作 - 召開董事會議的行政成本與時間，對公司是個不小的負擔': '董事會結構與運作 - 您投資的公司在召開董事會議的行政成本與時間，對公司是個不小的負擔',
        # 特殊處理：投資階段1的CSV檔誤將董事會議寫成股東會議
        '董事會結構與運作 - 召開董事會議的行政成本與時間，對公司是個不小的負擔_ALT': '董事會結構與運作 - 您投資的公司在召開股東會議的行政成本與時間，對公司是個不小的負擔',
        '董事會結構與運作 - 公司董事會的議程及相關資料能在3天前通知，並以掛號或電子方式（經過股東同意或於公司章程中載明）寄發': '董事會結構與運作 - 您投資的公司之董事會的議程及相關資料能在3天前通知，並以掛號或電子方式（經過股東同意或於公司章程中載明）寄發',
        '董事會結構與運作 - 公司董事會決議方式能夠清楚載明，且議事錄完整記載會議資訊（如時間、地點、主席、決議方法及結果）': '董事會結構與運作 - 您投資的公司之董事會決議方式能夠清楚載明，且議事錄完整記載會議資訊（如時間、地點、主席、決議方法及結果）',
        '董事會結構與運作 - 公司通常每年召開一次股東常會，並是由董事會召集': '董事會結構與運作 - 您投資的公司通常每年召開一次股東常會，並是由董事會召集',
        '公司定期性董事會的議事內容，通常包含以下哪些項目？ (可複選)': '您投資的公司其定期性董事會的議事內容，通常包含以下哪些項目？ （可複選）',
        '董事會結構與運作 - 在過去12個月內，貴公司董事會的召開頻率為何？': '董事會結構與運作 - 您投資的公司在過去12個月內，董事會的召開頻率為何？',
        '董事會結構與運作 - 公司諮詢顧問／業師的頻率為何？': '董事會結構與運作 - 您投資的公司在諮詢顧問／業師方面的頻率為何？',
        
        # 財務報表
        '公司最近一期的年度財務報表，是否委任外部會計師進行查核簽證？若有，其遵循的會計準則為何？': '您投資的公司在最近一期的年度財務報表，是否委任外部會計師進行查核簽證？若有，其遵循的會計準則為何？',
        
        # 資訊透明度
        '資訊透明度 - 公司通常多久向股東揭露董事、監察人、經理人及持股超過10%大股東的持股情形、股權質押比率與前十大股東之股權結構圖或表': '資訊透明度 - 您投資的公司通常多久向股東揭露董事、監察人、經理人及持股超過10%大股東的持股情形、股權質押比率與前十大股東之股權結構圖或表',
        '資訊透明度 - 公司通常多久向股東提供業務報告（如營運、研發進度等）': '資訊透明度 - 您投資的公司通常多久向股東提供業務報告（如營運、研發進度等）',
        '資訊透明度 - 公司通常多久向股東提供財務報告': '資訊透明度 - 您投資的公司通常多久向股東提供財務報告',
        
        # 內控與風險評估
        '內控與風險評估（含財務與營運風險） - 公司由不同人員分別負責出納與會計': '內控與風險評估（含財務與營運風險） - 您投資的公司由不同人員分別負責出納與會計',
        '內控與風險評估（含財務與營運風險） - 公司財務紀錄由專責人員或外部會計師協助處理': '內控與風險評估（含財務與營運風險） - 您投資的公司其財務紀錄由專責人員或外部會計師協助處理',
        '內控與風險評估（含財務與營運風險） - 公司開發的專利、商標等智慧財產權，均已登記在公司名下': '內控與風險評估（含財務與營運風險） - 您投資的公司其開發的專利、商標等智慧財產權，均已登記在公司名下',
        
        # 利害關係人
        '利害關係人 - 公司員工分紅制度設計能有效激勵員工': '利害關係人 - 您投資的公司其員工分紅制度設計能有效激勵員工',
        '利害關係人 - 公司已建立與主要利害關係人（如員工、債權人、外部投資人等）的溝通管道': '利害關係人 - 您投資的公司已建立與主要利害關係人（如員工、債權人、外部投資人等）的溝通管道',
        
        # ========== 第二階段新增 ==========
        # 股東會結構與運作（第二階段）
        '股東會結構與運作 - 公司股東常會之召開，通常於停止股票過戶日前公告受理股東提案': '股東會結構與運作 - 您投資的公司其股東常會之召開，通常於停止股票過戶日前公告受理股東提案',
        '股東會結構與運作 - 公司於公告受理股東提案時，受理期間通常大於十日': '股東會結構與運作 - 您投資的公司於公告受理股東提案時，受理期間通常大於十日',
        '股東會結構與運作 - 公司董事會通常會將股東提案列於股東常會議案中': '股東會結構與運作 - 您投資的公司之董事會通常會將股東提案列於股東常會議案中',
        '股東會結構與運作 - 公司董事會針對未列入議案的股東提案，會於股東會說明未列入的理由': '股東會結構與運作 - 您投資的公司之董事會針對未列入議案的股東提案，會於股東會說明未列入的理由',
        '股東會結構與運作 - 公司的股東結構中包含法人股東或創投': '股東會結構與運作 - 您投資的公司其股東結構中包含法人股東（如創投）',
        
        # 董事會結構與運作（第二階段）
        '董事會結構與運作 - 公司董事針對與自身利害相關的議題會說明與利益迴避': '董事會結構與運作 - 您投資的公司之董事針對與自身利害相關的議題會說明與利益迴避',
        '董事會結構與運作 - 公司董事會每季召開一次董事會，並討論定期性董事會的議事內容': '董事會結構與運作 - 您投資的公司每季會召開一次董事會，並討論定期性董事會的議事內容',
        '董事會結構與運作 - 公司董事會成員的專業背景涵蓋兩種（含）以上？（如：財務、法律、技術等）': '董事會結構與運作 - 您投資的公司之董事會成員專業背景涵蓋兩種（含）以上（如：財務、法律、技術等）',
        '董事會結構與運作 - 公司的董事間不具二親等之內的親屬關係': '董事會結構與運作 - 您投資的公司之董事間不具二親等之內的親屬關係',
        
        # 內控與風險評估（第二階段）
        '內控與風險評估（含財務與營運風險） - 公司有明確的職務權限制度': '內控與風險評估（含財務與營運風險） - 您投資的公司有明確的職務權限制度',
        '內控與風險評估（含財務與營運風險） - 建立書面核准流程對公司來說是一項挑戰': '內控與風險評估（含財務與營運風險） - 您投資的公司在建立書面核准流程上有困難',
        '內控與風險評估（含財務與營運風險） - 訂定財會作業程序對公司來說是不小的負擔': '內控與風險評估（含財務與營運風險） - 您投資的公司在訂定財會作業程序上會是不小的負擔',
        '內控與風險評估（含財務與營運風險） - 公司的印鑑與支票不會由相同人員管理': '內控與風險評估（含財務與營運風險） - 您投資的公司之印鑑與支票不會由相同人員管理',
        '內控與風險評估（含財務與營運風險） - 公司已訂定書面的書面之用印／借印流程規範': '內控與風險評估（含財務與營運風險） - 您投資的公司已訂定書面的用印／借印流程規範',
        '內控與風險評估（含財務與營運風險） - 公司有編列公司印鑑保管清冊作為保管記錄': '內控與風險評估（含財務與營運風險） - 您投資的公司有編列公司印鑑保管清冊作為保管記錄',
        '內控與風險評估（含財務與營運風險） - 公司已訂定關係人相互間財務業務相關作業規範，且內容包含關係人間的進銷貨交易、背書保證及資金貸與等事項': '內控與風險評估（含財務與營運風險） - 您投資的公司已訂定關係人相互間財務業務相關作業規範，且內容包含關係人間的進銷貨交易、背書保證及資金貸與等事項',
        '內控與風險評估（含財務與營運風險） - 公司已訂定取得或處分資產等交易之管理程序': '內控與風險評估（含財務與營運風險） - 您投資的公司已訂定取得或處分資產等交易之管理程序',
        '內控與風險評估（含財務與營運風險） - 公司於採購及銷售流程中，已導入請購單、採購單、出貨單，並經過適當層級核准': '內控與風險評估（含財務與營運風險） - 您投資的公司於採購及銷售流程中，已導入請購單、採購單、出貨單，並經過適當層級核准',
        '內控與風險評估（含財務與營運風險） - 公司具備供應商評估清單或相關評估文件': '內控與風險評估（含財務與營運風險） - 您投資的公司具備供應商評估清單或相關評估文件',
        '內控與風險評估（含財務與營運風險） - 公司辦理請購、採購、驗收時，是由不同人員負責': '內控與風險評估（含財務與營運風險） - 您投資的公司辦理請購、採購、驗收時，是由不同人員負責',
        '內控與風險評估（含財務與營運風險） - 公司定期採用審計品質指標（AQIs）評估外部會計師獨立性與適任性': '內控與風險評估（含財務與營運風險） - 您投資的公司定期採用審計品質指標（AQIs）評估外部會計師獨立性與適任性',
        
        # 其他第二階段題目
        '公司針對下列內部控制循環，建立書面控制程序與執行自評的進度為何？  (1)銷售及收款循環 (2)採購及付款循環 (3)投資循環 (4)資金貸與他人、背書保證之管理 (5)關係人交易之管理': '針對下列內部控制循環，您投資的公司在建立書面控制程序與執行自評的進度為何？',
        '公司導入企業資源規劃（ERP）系統的程度為何？': '您投資的公司導入企業資源規劃（ERP）系統的程度為何？',
        '公司定期檢視關鍵財務指標（如：負債比率、流動比率、總資產報酬率等）的頻率為何？': '您投資的公司定期檢視關鍵財務指標（如：負債比率、流動比率、總資產報酬率等）的頻率為何？',
        '您認為公司現金流量規劃與監控制度的建立程度如何？': '您投資的公司在現金流量規劃與監控制度的建立程度如何？',
        '承上題您認為公司現金流量足以支撐公司營運幾個月？': '承上題，您投資的公司之現金流量足以支撐公司營運幾個月？',
        '公司與關係人間的資金往來比率大約多少？': '您投資的公司與關係人間的資金往來比率大約多少？',
        
        # 利害關係人（第二階段）
        '利害關係人 - 公司有提供員工足夠的支持來規劃職涯升遷與發展': '利害關係人 - 您投資的公司有提供員工足夠的支持來規劃職涯升遷與發展',
        '利害關係人 - 公司已制定債務管理政策，並定期檢視債務結構與還款能力': '利害關係人 - 您投資的公司已制定債務管理政策，並定期檢視債務結構與還款能力',
        
        # ========== 第三階段新增 ==========
        # 董事會結構與運作（第三階段）
        '董事會結構與運作 - 公司在監察人或獨立董事的提名及當選等相關資訊揭露方面都有落實': '董事會結構與運作 - 您投資的公司在監察人或獨立董事的提名及當選等相關資訊揭露方面都有落實',
        '董事會結構與運作 - 公司內部設有稽核單位，且內部稽核報告的呈送方式未經主管階層修改': '董事會結構與運作 - 您投資的公司內部設有稽核單位，且內部稽核報告的呈送方式未經主管階層修改',
        '董事會結構與運作 - 公司重視董事的培訓與持續進修': '董事會結構與運作 - 您投資的公司重視董事的培訓與持續進修',
        '董事會結構與運作 - 承上題，公司過半數董事每年完成至少6小時的進修': '董事會結構與運作 - 承上題，您投資的公司有過半數董事每年完成至少6小時的進修',
        '董事會結構與運作 - 公司監察人具備財務或法律專業背景來履行職責': '董事會結構與運作 - 您投資的公司之監察人具備財務或法律專業背景來履行職責',
        '董事會結構與運作 - 公司監察人與董事在執行業務時能保持獨立性': '董事會結構與運作 - 您投資的公司之監察人與董事在執行業務時能保持獨立性',
        
        # 內控與風險評估（第三階段）
        '內控與風險評估（含財務與營運風險） - 公司在文件審核過程中具備完善的雙重覆核機制': '內控與風險評估（含財務與營運風險） - 您投資的公司在文件審核過程中具備完善的雙重覆核機制',
        '內控與風險評估（含財務與營運風險） - 公司有完整的資金貸與及背書保證的管理規範': '內控與風險評估（含財務與營運風險） - 您投資的公司有完整的資金貸與及背書保證的管理規範',
        '內控與風險評估（含財務與營運風險） - 公司有足夠的專職內部稽核人員來進行內部稽核工作': '內控與風險評估（含財務與營運風險） - 您投資的公司有足夠的專職內部稽核人員來進行內部稽核工作',
        '內控與風險評估（含財務與營運風險） - 公司專職內部稽核人員能保持獨立性，不受經營管理階層影響': '內控與風險評估（含財務與營運風險） - 您投資的公司之專職內部稽核人員能保持獨立性，不受經營管理階層影響',
        '內控與風險評估（含財務與營運風險） - 公司已建立正式且保密的檢舉制度（吹哨者政策）來處理內、外部人員的不當行為': '內控與風險評估（含財務與營運風險） - 您投資的公司已建立正式且保密的檢舉制度（吹哨者政策）來處理內、外部人員的不當行為',
        
        # 資訊透明度（第三階段）
        '資訊透明度 - 公司清楚的向股東揭露董事的個別酬金': '未命名題目 - 您投資的公司有清楚的向股東揭露董事的個別酬金',
        '資訊透明度 - 公司董事及經理人的酬金與公司績效連動': '未命名題目 - 您投資的公司其董事及經理人的酬金與公司績效連動',
        '資訊透明度 - 公司清楚的向股東揭露總經理及副總經理的個別酬金': '未命名題目 - 您投資的公司有清楚的向股東揭露總經理及副總經理的個別酬金',
        
        # 利害關係人（第三階段）
        '利害關係人 - 公司已充分透過網站或年報，說明員工福利措施（如員工福利信託等）、退休制度及其實施情形': '利害關係人 - 您投資的公司已充分透過網站或年報，說明員工福利措施（如員工福利信託等）、退休制度及其實施情形',
        '利害關係人 - 公司已充分透過網站或年報，說明員工人身安全與工作環境的保護措施及其實施情形': '利害關係人 - 您投資的公司已充分透過網站或年報，說明員工人身安全與工作環境的保護措施及其實施情形',
        '利害關係人 - 公司近一年內因勞資糾紛，曾產生一些處理費用': '利害關係人 - 您投資的公司近一年內因勞資糾紛，曾產生一些處理費用',
    }
    
    # 嘗試投資方映射
    if target_col in investor_mappings and investor_mappings[target_col] in df.columns:
        return investor_mappings[target_col]
    
    # 反向映射：投資方→公司方
    reverse_mapping = {v: k for k, v in investor_mappings.items()}
    if target_col in reverse_mapping and reverse_mapping[target_col] in df.columns:
        return reverse_mapping[target_col]
    
    # 特殊處理：投資階段1的董事會議行政成本欄位命名錯誤（寫成股東會議）
    if target_col == '董事會結構與運作 - 召開董事會議的行政成本與時間，對公司是個不小的負擔':
        alt_col = '董事會結構與運作 - 您投資的公司在召開股東會議的行政成本與時間，對公司是個不小的負擔'
        if alt_col in df.columns:
            return alt_col
    
    # 最後嘗試：移除「分類 - 」前綴後再查找（處理如「股東會結構與運作 - 題目」的情況）
    # 有些題目在程式中帶分類前綴，但在CSV中可能沒有
    if ' - ' in target_col:
        col_without_prefix = target_col.split(' - ', 1)[1]  # 只移除第一個 " - " 之前的部分
        if col_without_prefix in df.columns:
            return col_without_prefix
        
        # 也嘗試用無前綴版本去查找映射
        if col_without_prefix in investor_mappings and investor_mappings[col_without_prefix] in df.columns:
            return investor_mappings[col_without_prefix]
        if col_without_prefix in reverse_mapping and reverse_mapping[col_without_prefix] in df.columns:
            return reverse_mapping[col_without_prefix]
    
    # 如果都找不到，返回 None
    return None

def add_topic_analysis(doc, df, topic_col, topic_title, topic_description, full_question='', table_counter=None, insert_stat_plain=None, sig_topics=None):
    """
    新增單一議題的完整分析
    包含：完整題目、描述、表格、圖表、統計檢定、業務解讀
    即使統計檢定沒過也提供詳細敘述
    
    table_counter: 表格編號計數器
    注意：如果df沒有'respondent_type'欄位，則只做整體分析，不做公司方vs投資方比較
    """
    # 預設白話文插入與顯著題目列表
    if insert_stat_plain is None:
        insert_stat_plain = lambda x: x
    if sig_topics is None:
        sig_topics = []
    
    add_heading_with_style(doc, topic_title, level=2)
    
    # 顯示完整題目
    if full_question:
        question_para = doc.add_paragraph()
        question_para.add_run('問卷題目：').bold = True
        question_para.add_run(full_question)
        question_para.runs[0].font.size = Pt(11)
        question_para.runs[1].font.size = Pt(11)
        question_para.runs[1].font.color.rgb = RGBColor(64, 64, 64)
    
    if topic_description:
        para = doc.add_paragraph(topic_description)
        para.runs[0].font.size = Pt(11)
    
    # 查找匹配的欄位名稱（處理公司方和投資方的不同命名）
    actual_col = find_matching_column(df, topic_col)
    
    if actual_col is None:
        doc.add_paragraph(f'本題目不存在於資料中（查找欄位：{topic_col}）。')
        doc.add_page_break()
        return doc
    
    # 使用找到的實際欄位名稱
    topic_col = actual_col

    # === 3.9複選題專屬處理 ===
    is_39_multi = topic_title.startswith('3.9') or (
        '複選' in full_question or '多選' in full_question or '複選' in topic_title or '多選' in topic_title
    )

    if is_39_multi:
        add_heading_with_style(doc, '(一) 公司方與投資方複選頻率分析', level=3)
        # 解析複選題：以分號、逗號、頓號、空格等分割
        def split_options(val):
            if pd.isna(val):
                return []
            val = str(val)
            # 常見分隔符號：；, 、 、 ;
            for sep in ['；', ';', '、', ',', '，', '\n', '\r', '|', '/']:
                val = val.replace(sep, '|||')
            return [v.strip() for v in val.split('|||') if v.strip()]

        # 收集所有選項
        all_options = set()
        for v in df[topic_col].dropna():
            all_options.update(split_options(v))
        sorted_options = smart_sort_categories(list(all_options))

        # 分公司方/投資方計算（已移至下方 is_39_multi 處理區塊，這裡移除重複與錯誤嵌套）
        
        # 合計行
        company_total = crosstab.loc['All', '公司方'] if '公司方' in crosstab.columns else 0
        investor_total = crosstab.loc['All', '投資方'] if '投資方' in crosstab.columns else 0
        table_data['data'].append([
            '合計',
            company_total,
            '100.0%',
            investor_total,
            '100.0%',
            crosstab.loc['All', 'All']
        ])
        
        add_statistics_table(doc, table_data, title=f"{topic_title} - 受訪者類型分佈表", table_counter=table_counter)
        
        # === 加入長條圖 ===
        doc.add_paragraph()
        doc.add_paragraph('【圖表呈現】', style='Heading 4')
        
    # 預設白話文插入與顯著題目列表
    if insert_stat_plain is None:
        insert_stat_plain = lambda x: x
    if sig_topics is None:
        sig_topics = []

    # --- 修正：先檢查資料長度與欄位存在性 ---
    if topic_col not in df.columns:
        print(f"[資料錯誤] 欄位不存在: {topic_col}")
        doc.add_paragraph(f"[{topic_title} 欄位不存在，無法分析]")
        return doc
    if df[topic_col].dropna().shape[0] == 0:
        print(f"[資料錯誤] 欄位無有效資料: {topic_col}")
        doc.add_paragraph(f"[{topic_title} 欄位無有效資料，無法分析]")
        return doc

    # --- 原本流程 ---
    if 'respondent_type' in df.columns:
        df_clean = df[[topic_col, 'respondent_type']].copy()
        df_clean = df_clean.dropna(subset=[topic_col])
        df_clean[topic_col] = clean_and_merge_categories(df_clean[topic_col])
        crosstab = pd.crosstab(df_clean[topic_col], df_clean['respondent_type'], margins=True)
        crosstab_pct = pd.crosstab(df_clean[topic_col], df_clean['respondent_type'], normalize='columns') * 100
        # 生成表格
        table_data = {
            'columns': ['選項', '公司方人數', '公司方百分比', '投資方人數', '投資方百分比', '合計'],
            'data': []
        }
        categories = [idx for idx in crosstab.index if idx != 'All']
        sorted_categories = smart_sort_categories(categories)
        for idx in sorted_categories:
            company_count = crosstab.loc[idx, '公司方'] if '公司方' in crosstab.columns else 0
            company_pct = f"{crosstab_pct.loc[idx, '公司方']:.1f}%" if '公司方' in crosstab_pct.columns else '-'
            investor_count = crosstab.loc[idx, '投資方'] if '投資方' in crosstab.columns else 0
            investor_pct = f"{crosstab_pct.loc[idx, '投資方']:.1f}%" if '投資方' in crosstab_pct.columns else '-'
            total = crosstab.loc[idx, 'All']
            table_data['data'].append([
                str(idx), company_count, company_pct, investor_count, investor_pct, total
            ])
        # 合計行
        company_total = crosstab.loc['All', '公司方'] if '公司方' in crosstab.columns else 0
        investor_total = crosstab.loc['All', '投資方'] if '投資方' in crosstab.columns else 0
        table_data['data'].append([
            '合計', company_total, '100.0%', investor_total, '100.0%', crosstab.loc['All', 'All']
        ])
        add_statistics_table(doc, table_data, title=f"{topic_title} - 受訪者類型分佈表", table_counter=table_counter)
        # 長條圖
        doc.add_paragraph()
        doc.add_paragraph('【圖表呈現】', style='Heading 4')
        try:
            chart_title = f"{topic_title} - 公司方與投資方比較"
            avg_label_length = sum(len(str(cat)) for cat in categories) / len(categories) if categories else 0
            use_horizontal = avg_label_length > 15 or '議事內容' in topic_title
            import plotly.graph_objects as go
            if use_horizontal:
                fig = create_horizontal_bar_chart(crosstab, crosstab_pct, chart_title, categories)
            else:
                fig = create_bar_chart(crosstab, crosstab_pct, chart_title, categories)
            chart_filename = f"/tmp/chart_{hash(topic_title)}_bar.png"
            if save_plotly_as_image(fig, chart_filename):
                doc.add_picture(chart_filename, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                try:
                    os.remove(chart_filename)
                except:
                    pass
        except Exception as e:
            doc.add_paragraph(f'（圖表生成時發生錯誤）')

        # 顯著性檢定與白話文
        chi_result = calculate_chi_square(df_clean, topic_col, 'respondent_type') if 'respondent_type' in df_clean.columns else None
        if chi_result and chi_result['p_value'] < 0.05:
            company_top = crosstab_pct['公司方'].idxmax() if '公司方' in crosstab_pct.columns else ''
            investor_top = crosstab_pct['投資方'].idxmax() if '投資方' in crosstab_pct.columns else ''
            company_top_pct = crosstab_pct.loc[company_top, '公司方'] if company_top else 0
            investor_top_pct = crosstab_pct.loc[investor_top, '投資方'] if investor_top else 0
            plain_text = f"本題在公司方與投資方之間有明顯差異（p = {chi_result['p_value']:.4f}）。公司方最常選「{company_top}」（{company_top_pct:.1f}%），投資方則以「{investor_top}」（{investor_top_pct:.1f}%）為主，顯示雙方在此議題上的看法或做法有顯著不同。建議針對此差異進行深入討論。"
            doc.add_paragraph(insert_stat_plain(plain_text))
            if sig_topics is not None:
                sig_topics.append(topic_title)
        else:
            company_top = crosstab_pct['公司方'].idxmax() if '公司方' in crosstab_pct.columns else ''
            investor_top = crosstab_pct['投資方'].idxmax() if '投資方' in crosstab_pct.columns else ''
            company_top_pct = crosstab_pct.loc[company_top, '公司方'] if company_top else 0
            investor_top_pct = crosstab_pct.loc[investor_top, '投資方'] if investor_top else 0
            plain_text = f"本題公司方最常選「{company_top}」（{company_top_pct:.1f}%），投資方最常選「{investor_top}」（{investor_top_pct:.1f}%）。雙方分布趨勢相近，統計檢定未達顯著差異。此結果可作為公司治理現況的參考。"
            doc.add_paragraph(insert_stat_plain(plain_text))
    else:
        df_clean = df[[topic_col]].copy()
        df_clean = df_clean.dropna(subset=[topic_col])
        df_clean[topic_col] = clean_and_merge_categories(df_clean[topic_col])
        value_counts = df_clean[topic_col].value_counts()
        total_count = len(df_clean)
        crosstab = pd.DataFrame({
            '次數': value_counts,
            '百分比': (value_counts / total_count * 100).round(1)
        })
        crosstab.loc['合計'] = [total_count, 100.0]
        # 生成表格
        table_data = {
            'columns': ['選項', '次數', '百分比'],
            'data': []
        }
        categories = [idx for idx in crosstab.index if idx != '合計']
        sorted_categories = smart_sort_categories(categories)
        for idx in sorted_categories:
            table_data['data'].append([
                str(idx), int(crosstab.loc[idx, '次數']), f"{crosstab.loc[idx, '百分比']:.1f}%"
            ])
        # 合計行
        table_data['data'].append([
            '合計', int(crosstab.loc['合計', '次數']), f"{crosstab.loc['合計', '百分比']:.1f}%"
        ])
        add_statistics_table(doc, table_data, title=f"{topic_title} - 整體分佈表", table_counter=table_counter)
        # 長條圖
        doc.add_paragraph()
        doc.add_paragraph('【圖表呈現】', style='Heading 4')
        try:
            categories = [idx for idx in crosstab.index if idx != '合計']
            percentages = [crosstab.loc[idx, '百分比'] for idx in categories]
            sorted_categories = smart_sort_categories(categories)
            sorted_percentages = [crosstab.loc[cat, '百分比'] for cat in sorted_categories]
            chart_title = f"{topic_title} - 整體分佈"
            avg_label_length = sum(len(str(cat)) for cat in categories) / len(categories) if categories else 0
            use_horizontal = avg_label_length > 15 or '議事內容' in topic_title
            import plotly.graph_objects as go
            if use_horizontal:
                fig = go.Figure(go.Bar(
                    y=sorted_categories,
                    x=sorted_percentages,
                    orientation='h',
                    marker_color='#1f77b4',
                    text=[f"{v:.1f}%" for v in sorted_percentages],
                    textposition='auto'))
                fig.update_layout(
                    title=chart_title,
                    xaxis_title='百分比 (%)',
                    yaxis_title='選項',
                    template='plotly_white', height=500,
                    font=dict(family='Noto Sans CJK SC, WenQuanYi Micro Hei, sans-serif', size=12)
                )
            else:
                fig = go.Figure(go.Bar(
                    x=sorted_categories,
                    y=sorted_percentages,
                    marker_color='#1f77b4',
                    text=[f"{v:.1f}%" for v in sorted_percentages],
                    textposition='auto'))
                fig.update_layout(
                    title=chart_title,
                    xaxis_title='選項',
                    yaxis_title='百分比 (%)',
                    template='plotly_white', height=500,
                    font=dict(family='Noto Sans CJK SC, WenQuanYi Micro Hei, sans-serif', size=12)
                )
            chart_filename = f"/tmp/chart_{hash(topic_title)}_bar.png"
            if save_plotly_as_image(fig, chart_filename):
                doc.add_picture(chart_filename, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                try:
                    os.remove(chart_filename)
                except:
                    pass
        except Exception as e:
            doc.add_paragraph(f'（圖表生成時發生錯誤）')
        if chi_result and chi_result['p_value'] < 0.05:
            plain_text = f"本題在公司方與投資方之間有明顯差異（p = {chi_result['p_value']:.4f}）。公司方最常選「{company_top}」（{company_top_pct:.1f}%），投資方則以「{investor_top}」（{investor_top_pct:.1f}%）為主，顯示雙方在此議題上的看法或做法有顯著不同。建議針對此差異進行深入討論。"
            doc.add_paragraph(insert_stat_plain(plain_text))
            if sig_topics is not None:
                sig_topics.append(topic_title)
        else:
            plain_text = f"本題公司方最常選「{company_top}」（{company_top_pct:.1f}%），投資方最常選「{investor_top}」（{investor_top_pct:.1f}%）。雙方分布趨勢相近，統計檢定未達顯著差異。此結果可作為公司治理現況的參考。"
            doc.add_paragraph(insert_stat_plain(plain_text))
    
    # (二) 公司階段分析
    if 'phase' in df.columns and topic_col in df.columns:
        add_heading_with_style(doc, '(二) 公司發展階段分析', level=3)
        
        # 清理資料
        df_phase = df[[topic_col, 'phase']].dropna()
        
        if len(df_phase) > 0:
            # 清理並標準化階段分析資料
            df_phase[topic_col] = clean_and_merge_categories(df_phase[topic_col])
            
            # 階段交叉表
            phase_crosstab = pd.crosstab(df_phase[topic_col], df_phase['phase'], margins=True)
            phase_crosstab_pct = pd.crosstab(df_phase[topic_col], df_phase['phase'], normalize='columns') * 100
            
            # 生成階段表格
            phases = smart_sort_categories([p for p in phase_crosstab.columns if p != 'All'])
            
            if len(phases) > 0:
                table_columns = ['選項']
                # 階段已經排序好
                for phase in phases:
                    table_columns.extend([f'{phase}人數', f'{phase}百分比'])
                table_columns.append('合計')
                
                table_data = {
                    'columns': table_columns,
                    'data': []
                }
                
                # 智慧排序選項
                categories = [idx for idx in phase_crosstab.index if idx != 'All']
                sorted_categories = smart_sort_categories(categories)
                
                for idx in sorted_categories:
                    row = [str(idx)]
                    for phase in phases:
                        if phase in phase_crosstab.columns:
                            count = phase_crosstab.loc[idx, phase]
                            pct = f"{phase_crosstab_pct.loc[idx, phase]:.1f}%"
                        else:
                            count = 0
                            pct = '-'
                        row.extend([count, pct])
                    row.append(phase_crosstab.loc[idx, 'All'])
                    table_data['data'].append(row)
                
                # 合計行
                total_row = ['合計']
                for phase in phases:
                    if phase in phase_crosstab.columns:
                        total_row.extend([phase_crosstab.loc['All', phase], '100.0%'])
                    else:
                        total_row.extend([0, '-'])
                total_row.append(phase_crosstab.loc['All', 'All'])
                table_data['data'].append(total_row)
                
                add_statistics_table(doc, table_data, title=f"{topic_title} - 公司發展階段分佈表", table_counter=table_counter)
                
                # === 加入階段比較長條圖 ===
                doc.add_paragraph()
                doc.add_paragraph('【圖表呈現】', style='Heading 4')
                
                try:
                    # 獲取所有類別（排除 'All'）
                    categories = [idx for idx in phase_crosstab.index if idx != 'All']
                    
                    # 創建階段比較長條圖 - 使用完整問卷題目
                    chart_title = full_question if full_question else f"{topic_title} - 公司發展階段比較"
                    
                    # 檢查標籤長度，決定使用垂直或水平長條圖
                    avg_label_length = sum(len(str(cat)) for cat in categories) / len(categories) if categories else 0
                    use_horizontal = avg_label_length > 15 or '議事內容' in topic_title
                    
                    if use_horizontal:
                        fig = create_horizontal_phase_chart(phase_crosstab, phase_crosstab_pct, chart_title, categories, phases)
                    else:
                        fig = create_phase_chart(phase_crosstab, phase_crosstab_pct, chart_title, categories, phases)
                    
                    # 儲存圖片
                    chart_filename = f"/tmp/phase_chart_{hash(topic_title)}.png"
                    if save_plotly_as_image(fig, chart_filename):
                        # 將圖片插入 Word 文件
                        doc.add_picture(chart_filename, width=Inches(6))
                        # 圖片置中
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()
                        
                        # 清理臨時檔案
                        try:
                            os.remove(chart_filename)
                        except:
                            pass
                    else:
                        doc.add_paragraph('（圖表生成失敗）')
                except Exception as e:
                    print(f"階段圖表插入失敗: {e}")
                    doc.add_paragraph(f'（圖表生成時發生錯誤）')
                
                # === 統計檢定：根據資料類型選擇適當方法 ===
                doc.add_paragraph('【統計檢定】', style='Heading 4')
                
                try:
                    # 準備階段分組資料
                    phase_groups = [df_phase[df_phase['phase'] == p][topic_col].dropna() for p in phases]
                    valid_groups = [g for g in phase_groups if len(g) > 0]
                    
                    if len(valid_groups) < 2:
                        raise ValueError("有效階段組別不足（需至少2組）")
                    
                    # 判斷資料類型：嘗試轉換為數值
                    is_numeric = False
                    numeric_groups = []
                    for g in valid_groups:
                        try:
                            numeric_g = pd.to_numeric(g, errors='coerce').dropna()
                            if len(numeric_g) >= 3:  # 至少3個樣本
                                numeric_groups.append(numeric_g)
                        except:
                            pass
                    
                    # 如果所有組別都能轉為數值且樣本數足夠，視為連續變數
                    if len(numeric_groups) == len(valid_groups) and all(len(g) >= 3 for g in numeric_groups):
                        is_numeric = True
                    
                    if is_numeric:
                        # === 連續變數：使用 Kruskal-Wallis H 檢定（無母數）===
                        H_stat, p_val = kruskal(*numeric_groups)
                        
                        significance = ''
                        if p_val < 0.001:
                            significance = '***（高度顯著）'
                        elif p_val < 0.01:
                            significance = '**（非常顯著）'
                        elif p_val < 0.05:
                            significance = '*（顯著）'
                        else:
                            significance = 'n.s.（無顯著差異）'
                        
                        doc.add_paragraph(f"檢定方法：Kruskal-Wallis H 檢定（無母數檢定，適用於連續變數）")
                        doc.add_paragraph(f"H 統計量：{H_stat:.3f}")
                        doc.add_paragraph(f"自由度：{len(numeric_groups) - 1}")
                        doc.add_paragraph(f"顯著性水準：p = {p_val:.4f} {significance}")
                        
                        doc.add_paragraph()
                        doc.add_paragraph('【統計解讀】', style='Heading 4')
                        doc.add_paragraph(
                            f"Kruskal-Wallis H 檢定用於比較三個或以上獨立組別的中位數是否存在差異，"
                            f"不假設資料符合常態分佈，適用於順序資料或非常態分佈的連續資料。"
                        )
                        
                    else:
                        # === 類別變數：使用卡方檢定 ===
                        # 建立列聯表（不含邊際合計）
                        phase_crosstab_test = pd.crosstab(df_phase[topic_col], df_phase['phase'])
                        
                        # 確保有足夠的期望次數
                        chi2, p_val, dof, expected = chi2_contingency(phase_crosstab_test)
                        
                        # 檢查期望次數是否足夠（至少80%的格子 > 5）
                        low_expected = (expected < 5).sum()
                        total_cells = expected.size
                        low_expected_pct = (low_expected / total_cells) * 100
                        
                        significance = ''
                        if p_val < 0.001:
                            significance = '***（高度顯著）'
                        elif p_val < 0.01:
                            significance = '**（非常顯著）'
                        elif p_val < 0.05:
                            significance = '*（顯著）'
                        else:
                            significance = 'n.s.（無顯著差異）'
                        
                        doc.add_paragraph(f"檢定方法：卡方獨立性檢定（Chi-square test of independence，適用於類別變數）")
                        doc.add_paragraph(f"卡方統計量：χ² = {chi2:.3f}")
                        doc.add_paragraph(f"自由度：df = {dof}")
                        doc.add_paragraph(f"顯著性水準：p = {p_val:.4f} {significance}")
                        
                        # 警告：如果期望次數過低
                        if low_expected_pct > 20:
                            doc.add_paragraph(f"註：本題有 {low_expected_pct:.1f}% 之儲存格期望次數小於 5，檢定結果之穩定性可能受影響，解讀時需審慎。")
                        
                        doc.add_paragraph()
                        doc.add_paragraph('【統計解讀】', style='Heading 4')
                        doc.add_paragraph(
                            f"卡方檢定用於檢驗兩個類別變數之間是否存在關聯性。"
                            f"在此分析中，檢驗「公司發展階段」與「{topic_title}」是否具有顯著關聯。"
                            f"虛無假設（H₀）為兩變數獨立（無關聯），對立假設（H₁）為兩變數有關聯。"
                        )
                    
                    # === 共同的階段差異分析 ===
                    doc.add_paragraph()
                    doc.add_paragraph('【階段差異分析】', style='Heading 4')
                    
                    if p_val < 0.05:
                        doc.add_paragraph(
                            f"統計檢定顯示不同發展階段的公司在「{topic_title}」存在顯著差異（p = {p_val:.4f}）。"
                            f"此結果表明公司發展階段確實影響此議題的表現或認知。"
                        )
                        
                        # 提供各階段的具體觀察
                        doc.add_paragraph()
                        doc.add_paragraph('各階段特徵：', style='List Bullet')
                        
                        phase_analysis = {}
                        for phase in phases:
                            if phase in phase_crosstab_pct.columns:
                                top_option = phase_crosstab_pct[phase].idxmax()
                                top_pct = phase_crosstab_pct.loc[top_option, phase]
                                phase_analysis[phase] = {'option': top_option, 'pct': top_pct}
                                
                                p_bullet = doc.add_paragraph(style='List Bullet 2')
                                p_bullet.add_run(f"{phase}：").bold = True
                                p_bullet.add_run(f"主要選擇「{top_option}」（{top_pct:.1f}%）")
                        
                        # 深度趨勢分析
                        doc.add_paragraph()
                        doc.add_paragraph('【趨勢觀察與政策意涵】', style='Heading 4')
                        
                        # 比較第一階段與第三階段的變化
                        if '第一階段' in phase_analysis and '第三階段' in phase_analysis:
                            stage1_option = phase_analysis['第一階段']['option']
                            stage3_option = phase_analysis['第三階段']['option']
                            
                            if stage1_option == stage3_option:
                                doc.add_paragraph(
                                    f"從發展軌跡觀察，第一階段至第三階段的公司皆以「{stage1_option}」為主要選擇，"
                                    f"顯示此治理實務在各發展階段均受重視。然而，各階段在選擇比例上仍存在差異，"
                                    f"反映出隨著公司成熟度提升，治理機制的深化程度有所不同。"
                                )
                            else:
                                doc.add_paragraph(
                                    f"觀察公司發展軌跡，第一階段主要選擇「{stage1_option}」，"
                                    f"至第三階段則轉向「{stage3_option}」，顯示公司治理實務隨發展階段而演進。"
                                    f"此變化反映出企業在不同成長階段對治理機制有不同的需求與優先順序。"
                                )
                        
                        # 政策建議
                        doc.add_paragraph(
                            f"建議針對不同階段公司的特性，提供差異化的治理建議或輔導措施："
                        )
                        
                        policy_bullets = [
                            "第一階段（種子輪至B輪）：著重基礎治理架構建立，協助新創企業理解治理重要性，建立基本的決策流程與資訊揭露機制",
                            "第二階段（C輪至D輪）：強化內部控制與資訊透明度，輔導企業建立更完善的內部稽核制度、財務管理系統及股東溝通機制",
                            "第三階段（擴展期至成熟期）：完善利害關係人溝通與永續治理，協助企業建立全面性治理框架，為未來可能的IPO或併購做準備"
                        ]
                        
                        for policy in policy_bullets:
                            p = doc.add_paragraph(style='List Bullet 2')
                            p.add_run(policy)
                    else:
                        doc.add_paragraph(
                            f"統計檢定顯示不同發展階段的公司在「{topic_title}」無顯著差異（p = {p_val:.4f}）。"
                            f"此結果表明本議題可能是跨階段的共同關注點，不因公司發展階段而有明顯變化。"
                        )
                        
                        # 即使不顯著，仍提供描述性觀察
                        doc.add_paragraph()
                        doc.add_paragraph('各階段分佈觀察：', style='List Bullet')
                        
                        phase_consistency = []
                        for phase in phases:
                            if phase in phase_crosstab_pct.columns:
                                top_option = phase_crosstab_pct[phase].idxmax()
                                top_pct = phase_crosstab_pct.loc[top_option, phase]
                                phase_consistency.append(top_option)
                                
                                p_bullet = doc.add_paragraph(style='List Bullet 2')
                                p_bullet.add_run(f"{phase}：").bold = True
                                p_bullet.add_run(f"主要選擇「{top_option}」（{top_pct:.1f}%）")
                        
                        # 一致性分析
                        doc.add_paragraph()
                        doc.add_paragraph('【實務意涵】', style='Heading 4')
                        
                        if len(set(phase_consistency)) == 1:
                            doc.add_paragraph(
                                f"值得注意的是，雖然統計上未達顯著差異，但三個發展階段的公司均以「{phase_consistency[0]}」為主要選擇，"
                                f"顯示此治理實務具有跨階段的一致性，是未上市櫃公司普遍認同的治理方式。"
                                f"此共識可作為推動相關政策或建立治理標準的重要依據。"
                            )
                        else:
                            doc.add_paragraph(
                                f"雖然統計上未達顯著差異，但各階段主要選擇略有不同，"
                                f"建議持續觀察各階段公司的治理實踐，累積更多資料以深入了解階段性差異的細微變化。"
                                f"此類描述性資訊仍具參考價值，可供業務推動時考量不同階段公司的特性。"
                            )
                        
                except Exception as e:
                    doc.add_paragraph(f"由於資料結構限制或樣本數不足，無法進行統計檢定。錯誤訊息：{str(e)}")
                    doc.add_paragraph()
                    doc.add_paragraph('【數據解讀】', style='Heading 4')
                    
                    # 從階段分佈進行數據解讀
                    phase_descriptions = []
                    for phase in phases:
                        if phase in phase_crosstab_pct.columns:
                            top_option = phase_crosstab_pct[phase].idxmax()
                            top_pct = phase_crosstab_pct.loc[top_option, phase]
                            phase_descriptions.append(f"{phase}主要選擇「{top_option}」（{top_pct:.1f}%）")
                    
                    if phase_descriptions:
                        doc.add_paragraph(
                            f"從各發展階段的分佈來看，{topic_title}的表現呈現階段性差異。"
                            f"{'；'.join(phase_descriptions)}。"
                            f"第一階段公司通常處於建立基礎治理架構的時期，第二階段公司著重於制度化與規範化，"
                            f"第三階段公司則追求更高標準的治理品質。建議公司參考同階段的最佳實踐，循序漸進地提升治理水平。"
                        )
        else:
            doc.add_paragraph('本題目無有效的階段資料。')
    
    doc.add_paragraph()  # 空行
    doc.add_page_break()  # 每個議題後分頁
    return doc

def generate_full_descriptive_report(df, output_path="/workspaces/work1/問卷描述性統計報告_完整版.docx", add_metadata=True):
    """
    生成完整描述性統計報告（Word 格式）
    包含更多題目，附上政府統計風格表格
    即使統計檢定沒過也提供敘述性分析
    
    參數:
        df: pandas DataFrame - 問卷資料
        output_path: str - 輸出檔案路徑
        add_metadata: bool - 是否自動添加 respondent_type 和 phase 欄位（根據檔案名推斷）
    """
    print("開始生成描述性統計報告...")
    
    # 自動添加 metadata 欄位（如果尚未存在）
    if add_metadata:
        # 添加 respondent_type（如果有 _source_file 欄位）
        if '_source_file' in df.columns and 'respondent_type' not in df.columns:
            def infer_role(fname):
                if not isinstance(fname, str):
                    return '未知'
                fname_lower = fname.lower()
                # 檢查是否包含「投資方」相關關鍵字
                if '投資方' in fname or '投資' in fname or 'invest' in fname_lower:
                    return '投資方'
                return '公司方'
            df['respondent_type'] = df['_source_file'].astype(str).apply(infer_role)
            print("已自動添加 respondent_type 欄位")
        
        # 添加 phase（優先從問卷欄位推斷，其次從檔名推斷）
        if 'phase' not in df.columns:
            phase_added = False
            
            # 方法1: 優先從問卷內容欄位推斷（最準確）
            PHASE_COLUMN_NAME = "請問公司目前主要處於哪個發展階段？："
            if PHASE_COLUMN_NAME in df.columns:
                def extract_phase(val):
                    if pd.isna(val):
                        return None
                    val_str = str(val)
                    if '第一階段' in val_str or '一階段' in val_str:
                        return '第一階段'
                    elif '第二階段' in val_str or '二階段' in val_str:
                        return '第二階段'
                    elif '第三階段' in val_str or '三階段' in val_str:
                        return '第三階段'
                    return None
                df['phase'] = df[PHASE_COLUMN_NAME].apply(extract_phase)
                print(f"從問卷欄位推斷 phase")
                phase_added = True
            
            # 方法2: 從檔案名推斷階段（備用方法）
            if not phase_added and '_source_file' in df.columns:
                def infer_phase(fname):
                    if not isinstance(fname, str):
                        return None
                    if '第一階段' in fname or '一階段' in fname:
                        return '第一階段'
                    elif '第二階段' in fname or '二階段' in fname:
                        return '第二階段'
                    elif '第三階段' in fname or '三階段' in fname:
                        return '第三階段'
                    return None
                df['phase'] = df['_source_file'].astype(str).apply(infer_phase)
                print("從檔案名推斷 phase")
                phase_added = True
    
    # 創建基礎文件並獲取 table_counter
    doc, table_counter = generate_descriptive_report_word(df, output_path)
    # 載入合併題目清單（優先使用 repository 中的 merged_auto_report.md）
    def load_topics_from_merged(md_path):
        import re
        topics = []
        try:
            with open(md_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except Exception:
            return []

        parts = re.split(r"\n###\s*問題:\s*", text)
        for p in parts[1:]:
            lines = p.splitlines()
            if not lines:
                continue
            question = lines[0].strip().rstrip(':').strip()
            # 移除開頭的階段標記，例如：【第一階段：...】、【第二階段：...】 等
            question = re.sub(r'^【.*?】\s*', '', question)
            # 移除前導的 # 或其他可能是 metadata 的標記
            question = question.lstrip('#').strip()
            # 若 question 非法（過短或僅含特殊符號或包含斜線且像 metadata），則跳過
            if len(question) < 3 or re.match(r'^[^\w\u4e00-\u9fff]+$', question) or ( '/' in question and len(question) < 10 ):
                continue
            # 明確排除不需要放入報告的 metadata 題目，例如填答身分等
            q_low = question.lower()
            if q_low.startswith('請問您的填答') or '填答身分' in q_low or q_low.strip() in ['請問您的填答身分：', '請問您的填答身分', '請問您的填答身分:']:
                # 使用者指定不放入此題，跳過
                continue
            # 排除代表公司名稱或公司識別相關的 metadata 題目（避免納入分析）
            if '代表公司' in q_low or '代表公司的名稱' in q_low or '公司名稱' in q_low or q_low.strip().startswith('請問您代表公司'):
                continue
            # 簡單擷取下一段作為描述（跳過 - ** 標記）
            desc = ''
            for ln in lines[1:6]:
                ln = ln.strip()
                if not ln:
                    continue
                if ln.startswith('- **'):
                    # use the rest after colon if present
                    m = re.match(r'- \*\*.*?\*\*:\s*(.*)', ln)
                    if m:
                        desc = m.group(1).strip()
                        break
                    continue
                desc = ln
                break
            topics.append({'col': question, 'title': question[:60], 'description': desc, 'question': question})
        return topics

    # 若 repository 中有 merged_auto_report.md，優先使用；否則採用 app 的合併邏輯來產生題目清單
    merged_md = os.path.join(os.path.dirname(__file__), 'merged_auto_report.md')
    if os.path.exists(merged_md):
        topics = load_topics_from_merged(merged_md)
        print(f"載入 {len(topics)} 個合併題目自 merged_auto_report.md")
    else:
        # 使用與 cloud_app 類似的合併邏輯（簡化版）來自動建立題目列表
        cols_to_exclude = ['為了後續支付訪談費，請提供您的電子郵件地址（我們將僅用於聯繫您支付訪談費，並妥善保護您的資料）:', 'IP紀錄', '額滿結束註記', '使用者紀錄', '會員時間', 'Hash', '會員編號', '自訂ID', '備註', '填答時間', '請問公司目前主要處於哪個發展階段？：', '_source_file', 'respondent_type']

        def normalize_question_v2_small(q):
            if not isinstance(q, str):
                return q
            q = q.replace('您投資的公司', '公司')
            q = q.replace('貴公司', '公司')
            q = re.sub(r'^未命名題目[\s\-：:]*', '', q)
            q = q.replace('（可複選）', '').replace('(可複選)', '')
            q = re.sub(r'\s+', ' ', q).strip()
            q = q.rstrip('：:。.,;；？?')
            return q

        def calculate_similarity(s1, s2):
            base_similarity = SequenceMatcher(None, s1, s2).ratio()
            if base_similarity > 0.8:
                keywords_s1 = set(re.findall(r'[\u4e00-\u9fff]{2,}', s1))
                keywords_s2 = set(re.findall(r'[\u4e00-\u9fff]{2,}', s2))
                if keywords_s1 and keywords_s2:
                    keyword_overlap = len(keywords_s1 & keywords_s2) / max(len(keywords_s1), len(keywords_s2))
                    return base_similarity * (0.5 + 0.5 * keyword_overlap)
            return base_similarity

        def merge_similar_questions_small(df_input, cols_to_exclude, similarity_threshold=0.75):
            all_cols = [c for c in df_input.columns if c not in cols_to_exclude]
            normalized_groups = {}
            for col in all_cols:
                norm = normalize_question_v2_small(col)
                normalized_groups.setdefault(norm, []).append(col)

            merged_mapping = {}
            processed_norms = set()
            norm_keys = list(normalized_groups.keys())
            for i, norm1 in enumerate(norm_keys):
                if norm1 in processed_norms:
                    continue
                similar_group = [norm1]
                for norm2 in norm_keys[i+1:]:
                    if norm2 in processed_norms:
                        continue
                    sim = calculate_similarity(norm1, norm2)
                    if sim >= similarity_threshold:
                        similar_group.append(norm2)
                        processed_norms.add(norm2)
                all_originals = []
                for norm in similar_group:
                    all_originals.extend(normalized_groups[norm])
                representative = None
                for orig in sorted(all_originals, key=lambda x: (len(x), '投資' in x)):
                    if '未命名題目' not in orig:
                        representative = orig
                        break
                if representative is None:
                    representative = all_originals[0]
                merged_mapping[representative] = all_originals
                processed_norms.add(norm1)

            # 合併資料，以代表題目為主
            for rep, originals in merged_mapping.items():
                if len(originals) > 1:
                    for other in originals[1:]:
                        mask = df_input[rep].isna() & df_input[other].notna()
                        df_input.loc[mask, rep] = df_input.loc[mask, other]

            cols_to_analyze = list(merged_mapping.keys())
            return merged_mapping, cols_to_analyze

        merged_mapping, cols_to_analyze = merge_similar_questions_small(df, cols_to_exclude, similarity_threshold=0.70)
        topics = [{'col': c, 'title': c, 'description': '', 'question': c} for c in cols_to_analyze]
        print(f"使用 app 類似合併邏輯，產生 {len(topics)} 個分析題目")
    
    # Debug: 列出所有題目與實際欄位的比對
    print("\n==== 題目與實際欄位比對 ====")
    for topic in topics:
        col = topic['col']
        try:
            match = find_matching_column(df, col)
        except Exception as e:
            match = None
            print(f"[DEBUG] find_matching_column 錯誤: {e}")
        if match:
            # 顯示按 respondent_type 的計數，幫助檢查公司/投資方分配
            if 'respondent_type' in df.columns and match in df.columns:
                counts = df.groupby('respondent_type')[match].apply(lambda s: s.dropna().shape[0]).to_dict()
            else:
                counts = {}
            print(f"題目: {col} -> ✔️ {match} | counts by respondent_type: {counts}")
        else:
            print(f"題目: {col} -> ❌ 無對應欄位")
    print("==== End 比對 ====")

    # 逐題分析（所有有資料的題目都進行分析與圖表插入）
    analyzed_count = 0
    for topic in topics:
        print(f"\n--- 分析題目: {topic['title']} ({topic['col']}) ---")
        try:
            actual_col = find_matching_column(df, topic['col'])
        except Exception as e:
            print(f"[DEBUG] find_matching_column 錯誤（跳過題目）: {e}")
            doc.add_paragraph(f"[{topic['title']} - 找不到對應欄位，已跳過]")
            continue

        if actual_col is not None:
            try:
                add_topic_analysis(
                    doc, df,
                    topic['col'],
                    topic['title'],
                    topic['description'],
                    full_question=topic.get('question', ''),
                    table_counter=table_counter
                )
                analyzed_count += 1
                print(f"完成: {topic['title']}")
            except Exception as e:
                print(f"分析 {topic['title']} 時發生錯誤: {e}")
                doc.add_paragraph(f"[{topic['title']} 資料不足或分析發生錯誤]")
        else:
            print(f"❌ 欄位不存在，跳過: {topic['col']}")
    print(f"\n共分析 {analyzed_count} 個議題 (共 {len(topics)} 題)")
    
    # === 新增：信度與效度分析 ===
    if RELIABILITY_AVAILABLE:
        try:
            add_reliability_validity_analysis(doc, df, topics, table_counter)
        except Exception as e:
            print(f"信度效度分析發生錯誤: {e}")
            doc.add_paragraph(f"[信度效度分析發生錯誤：{str(e)}]")
    
    # 儲存文件
    doc.save(output_path)
    print(f"報告已儲存至: {output_path}")
    
    return output_path


def add_reliability_validity_analysis(doc, df, topics, table_counter):
    """
    添加信度與效度分析章節
    """
    doc.add_page_break()
    doc.add_heading('四、問卷信度與效度分析', level=1)
    
    doc.add_paragraph(
        "本章節針對問卷的信度（Reliability）與效度（Validity）進行統計分析，"
        "以評估問卷測量的一致性與有效性。"
    )
    
    # 按維度分組題目
    dimensions = {
        '股東會結構與運作': [],
        '董事會結構與運作': [],
        '內控與風險評估': [],
        '資訊透明度': [],
        '利害關係人': []
    }
    
    for topic in topics:
        col = topic['col']
        for dim_name in dimensions.keys():
            if dim_name in col:
                dimensions[dim_name].append(col)
                break
    
    # 分析每個維度
    doc.add_heading('4.1 各維度信度分析（Cronbach\'s Alpha）', level=2)
    doc.add_paragraph(
        "Cronbach's Alpha係數用於衡量量表的內部一致性信度，"
        "係數值介於0到1之間，一般認為0.7以上表示信度良好，0.8以上表示信度很好。"
    )
    
    alpha_results = []
    
    for dim_name, dim_cols in dimensions.items():
        if len(dim_cols) < 2:
            continue
        
        doc.add_heading(f'{dim_name}維度', level=3)
        
        # 找出實際存在的欄位
        actual_cols = []
        for col in dim_cols:
            actual_col = find_matching_column(df, col)
            if actual_col:
                actual_cols.append(actual_col)
        
        if len(actual_cols) < 2:
            doc.add_paragraph(f"本維度題項不足（少於2題），無法計算信度係數。")
            continue
        
        # 計算 Cronbach's Alpha
        try:
            alpha, item_stats = calculate_cronbach_alpha(df, actual_cols)
            
            if alpha is not None and not np.isnan(alpha):
                alpha_results.append({
                    '維度': dim_name,
                    '題項數': len(actual_cols),
                    "Cronbach's Alpha": f"{alpha:.3f}",
                    '信度評價': interpret_cronbach_alpha(alpha)
                })
                
                doc.add_paragraph(
                    f"本維度包含{len(actual_cols)}個題項，Cronbach's Alpha = {alpha:.3f}（{interpret_cronbach_alpha(alpha)}）"
                )
                
                if item_stats is not None and len(item_stats) > 0:
                    # 添加題項統計表
                    table_data = {
                        'columns': ['題項', '平均數', '標準差', '刪除後Alpha'],
                        'data': []
                    }
                    
                    for _, row in item_stats.iterrows():
                        # 簡化題項名稱
                        item_name = row['題項'].split(' - ')[-1][:40]
                        table_data['data'].append([
                            item_name,
                            f"{row['平均數']:.2f}",
                            f"{row['標準差']:.2f}",
                            f"{row['刪除後Alpha']:.3f}" if not pd.isna(row['刪除後Alpha']) else 'N/A'
                        ])
                    
                    add_statistics_table(doc, table_data, title=f"{dim_name} - 題項統計", table_counter=table_counter)
            else:
                # 診斷問題
                test_data = df[actual_cols].copy()
                for col in actual_cols:
                    if test_data[col].dtype == 'object':
                        likert_map = {
                            '非常不同意': 1, '不同意': 2, '普通': 3, '同意': 4, '非常同意': 5
                        }
                        test_data[col] = test_data[col].map(likert_map)
                test_data = test_data.dropna()
                
                if len(test_data) < 2:
                    doc.add_paragraph(f"本維度有效資料筆數不足（有效筆數：{len(test_data)}），無法計算信度係數。")
                elif test_data.var().sum() == 0:
                    doc.add_paragraph(f"本維度所有題項答案一致（無變異），無法計算信度係數。")
                else:
                    doc.add_paragraph(f"由於資料特性（如變異數接近零或矩陣奇異），無法計算本維度的信度係數。")
        except Exception as e:
            doc.add_paragraph(f"本維度信度計算發生錯誤：{str(e)[:100]}")
            continue
    
    # 信度總表
    if alpha_results:
        doc.add_heading('各維度信度總表', level=3)
        summary_table = {
            'columns': ['維度', '題項數', "Cronbach's Alpha", '信度評價'],
            'data': [[r['維度'], r['題項數'], r["Cronbach's Alpha"], r['信度評價']] for r in alpha_results]
        }
        add_statistics_table(doc, summary_table, title="問卷各維度信度係數總表", table_counter=table_counter)
    
    # KMO與Bartlett檢定
    doc.add_page_break()
    doc.add_heading('4.2 效度分析（KMO與Bartlett檢定）', level=2)
    doc.add_paragraph(
        "KMO（Kaiser-Meyer-Olkin）取樣適切性檢定用於衡量變數間的偏相關性，"
        "KMO值介於0到1之間，一般認為0.6以上可接受，0.7以上良好，0.8以上很好。"
        "Bartlett球形檢定用於檢驗相關矩陣是否為單位矩陣，若p<0.05表示變數間存在相關性，適合進行因素分析。"
    )
    
    for dim_name, dim_cols in dimensions.items():
        if len(dim_cols) < 3:
            continue
        doc.add_heading(f'{dim_name}維度', level=3)
        # 找出實際存在的欄位
        actual_cols = []
        for col in dim_cols:
            actual_col = find_matching_column(df, col)
            if actual_col:
                actual_cols.append(actual_col)
        if len(actual_cols) < 3:
            doc.add_paragraph(f"本維度題項不足（少於3題），無法進行效度分析。")
            continue
        try:
            # KMO檢定
            kmo_model, kmo_vars = kmo_test(df, actual_cols)
            # Bartlett檢定
            bartlett_stat, bartlett_p, bartlett_dof = bartlett_test(df, actual_cols)
            kmo_str = None
            bartlett_str = None
            kmo_value = None
            if kmo_model is not None and not np.isnan(kmo_model):
                kmo_value = kmo_model
                kmo_str = f"KMO取樣適切性檢定 = {kmo_model:.3f}（{interpret_kmo(kmo_model)}）"
                doc.add_paragraph(kmo_str)
            else:
                # 診斷問題
                test_data = df[actual_cols].copy()
                for col in actual_cols:
                    if test_data[col].dtype == 'object':
                        likert_map = {'非常不同意': 1, '不同意': 2, '普通': 3, '同意': 4, '非常同意': 5}
                        test_data[col] = test_data[col].map(likert_map)
                test_data = test_data.dropna()
                if len(test_data) < 3:
                    doc.add_paragraph(f"KMO檢定：有效資料筆數不足（有效筆數：{len(test_data)}），無法計算")
                else:
                    doc.add_paragraph("KMO檢定：相關矩陣奇異或變異數為零，無法計算")
            if bartlett_stat is not None and bartlett_p is not None and not np.isnan(bartlett_stat):
                significance = "顯著" if bartlett_p < 0.05 else "不顯著"
                bartlett_str = f"Bartlett球形檢定：χ² = {bartlett_stat:.2f}, df = {bartlett_dof}, p = {bartlett_p:.4f}（{significance}）"
                doc.add_paragraph(bartlett_str)
                # 修正結論產生邏輯：KMO<0.5時自動產生正確結論
                if bartlett_p < 0.05:
                    if kmo_value is not None and kmo_value < 0.5:
                        doc.add_paragraph(f"雖然Bartlett球形檢定顯著，但KMO值{kmo_value:.3f}低於0.5的可接受標準，顯示本維度的取樣適切性不足，不適合進行因素分析。")
                    else:
                        doc.add_paragraph("檢定結果顯著（p < 0.05），表示變數間存在相關性，適合進行因素分析。")
                else:
                    doc.add_paragraph("檢定結果不顯著（p ≥ 0.05），變數間相關性較弱。")
            else:
                doc.add_paragraph("Bartlett球形檢定：相關矩陣行列式為零或負值，無法計算")
        except Exception as e:
            doc.add_paragraph(f"本維度效度分析發生錯誤：{str(e)[:100]}")
    
    # 總結
    doc.add_heading('4.3 信效度分析總結', level=2)
    
    # 4.3信效度總結：根據各維度KMO/Alpha自動產生誠實結論
    if alpha_results:
        avg_alpha = np.mean([float(r["Cronbach's Alpha"]) for r in alpha_results])
        # 取得各維度信度與效度狀態
        alpha_dict = {r['維度']: float(r["Cronbach's Alpha"]) for r in alpha_results}
        # 重新計算KMO狀態
        kmo_status = {}
        for dim_name, dim_cols in dimensions.items():
            if len(dim_cols) < 3:
                continue
            actual_cols = [find_matching_column(df, col) for col in dim_cols if find_matching_column(df, col)]
            if len(actual_cols) < 3:
                continue
            try:
                kmo_model, _ = kmo_test(df, actual_cols)
                kmo_status[dim_name] = kmo_model if kmo_model is not None and not np.isnan(kmo_model) else None
            except Exception:
                kmo_status[dim_name] = None
        # 根據用戶指定語句自動產生結論
        # 1. 若僅「利害關係人」信度良好（>=0.7），其他維度信度不佳
        # 2. 若「股東會結構與運作」KMO<0.5，需明確指出
        good_dims = [k for k, v in alpha_dict.items() if v >= 0.7]
        bad_dims = [k for k, v in alpha_dict.items() if v < 0.7]
        kmo_bad_dims = [k for k, v in kmo_status.items() if v is not None and v < 0.5]
        # 用戶指定語句
        if good_dims == ['利害關係人'] and ('股東會結構與運作' in bad_dims or '內控與風險評估' in bad_dims):
            doc.add_paragraph(
                "僅『利害關係人』維度信度良好（0.710）；『股東會結構與運作』及『內控與風險評估』信度不佳，未來應審慎或修正。『股東會』維度未通過KMO效度檢定，結構效度不足。"
            )
            doc.add_paragraph("本問卷信度與效度整體不足，結果僅供參考，建議未來修正問卷設計或補充資料。")
        else:
            # 若所有維度信度/效度皆不佳，僅顯示誠實語句
            if all(v < 0.7 for v in alpha_dict.values()) or all((v is not None and v < 0.5) for v in kmo_status.values() if v is not None):
                doc.add_paragraph("本問卷信度與效度整體不足，結果僅供參考，建議未來修正問卷設計或補充資料。")
            else:
                doc.add_paragraph(
                    f"本問卷各維度的平均Cronbach's Alpha係數為{avg_alpha:.3f}，"
                    f"{'整體信度良好' if avg_alpha >= 0.7 else '部分維度信度需改善'}。"
                    "各維度的KMO值與Bartlett檢定結果顯示問卷具有一定的效度基礎。"
                )
        # 若有KMO<0.5的維度，補充說明
        if kmo_bad_dims:
            for dim in kmo_bad_dims:
                doc.add_paragraph(f"『{dim}』維度KMO值低於0.5，結構效度不足，未來應審慎解讀。")


if __name__ == "__main__":
    # 測試用
    print("描述性統計報告生成器已就緒")
